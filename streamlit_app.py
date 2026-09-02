from __future__ import annotations

import csv
import base64
import difflib
import html
import json
import io
import logging
import math
import os
import re
import smtplib
import urllib.parse
import urllib.request
import urllib.error
import uuid
import time as time_module
import zipfile
from collections.abc import Mapping
from pathlib import Path
from datetime import date, datetime, time, timedelta
from email.message import EmailMessage
from typing import Any
from zoneinfo import ZoneInfo

import gspread
import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


APP_TITLE = "폭염대비 온열질환 예방을 위한 조치사항"
APP_VERSION = "Professional UI v3.62 · 2026-09-02"
WORKSHEET_DEFAULT = "records"
SPREADSHEET_URL_FALLBACK = (
    "https://docs.google.com/spreadsheets/d/"
    "18c-qnfPmGG25qyAM497R7czDw3F7J7WRKmdLX3IGtY0"
)
KST = ZoneInfo("Asia/Seoul")
KMA_POINT_API_URL = (
    "https://apihub.kma.go.kr/api/typ01/cgi-bin/url/"
    "nph-sfc_obs_nc_pt_api"
)
KMA_ULTRA_NOWCAST_API_URL = (
    "https://apihub.kma.go.kr/api/typ02/openApi/"
    "VilageFcstInfoService_2.0/getUltraSrtNcst"
)
KMA_VILLAGE_FORECAST_API_URL = (
    "https://apihub.kma.go.kr/api/typ02/openApi/"
    "VilageFcstInfoService_2.0/getVilageFcst"
)
OPEN_METEO_ARCHIVE_API_URL = "https://archive-api.open-meteo.com/v1/archive"
KAKAO_PLACE_API_URL = "https://dapi.kakao.com/v2/local/search/keyword.json"
FORECAST_HEAT_THRESHOLD = 31.0
KMA_POINT_RANGE_TIMEOUT_SECONDS = 15
KMA_POINT_RANGE_MAX_CONSECUTIVE_TIMEOUTS = 3
KMA_REGIONAL_TIMEOUT_SECONDS = 12
KMA_REGIONAL_MAX_INITIAL_FAILURES = 3
OPEN_METEO_ARCHIVE_TIMEOUT_SECONDS = 15

OLD_CURRENT_COLUMNS = [
    "id",
    "작업날짜",
    "종목",
    "현장명",
    "팀",
    "근무시작",
    "근무종료",
    "작성자",
    "폭염시작",
    "폭염종료",
    "체감온도",
    "휴게시간",
    "공통 조치사항",
    "조치사항",
    "특이사항",
    "등록시간",
    "수정시간",
]

COLUMNS = [
    "id",
    "작업날짜",
    "종목",
    "현장명",
    "팀",
    "근무자수",
    "직원",
    "도급",
    "근무시작",
    "근무종료",
    "작성자",
    "폭염시작",
    "폭염종료",
    "체감온도",
    "휴게시간",
    "공통 조치사항",
    "조치사항",
    "특이사항",
    "등록시간",
    "수정시간",
]

# v3.16 시트 구조: 공통 조치사항이 조치사항에 합쳐져 있던 상태
PRE_COMMON_COLUMNS = [
    "id",
    "작업날짜",
    "종목",
    "현장명",
    "팀",
    "근무시작",
    "근무종료",
    "작성자",
    "폭염시작",
    "폭염종료",
    "체감온도",
    "휴게시간",
    "조치사항",
    "특이사항",
    "등록시간",
    "수정시간",
]

LEGACY_COLUMNS = [
    "id",
    "작업날짜",
    "현장명",
    "팀",
    "근무시작",
    "근무종료",
    "작성자",
    "작업인원",
    "폭염시작",
    "폭염종료",
    "체감온도",
    "휴게시작",
    "휴게종료",
    "휴게시간",
    "조치사항",
    "특이사항",
    "등록시간",
    "수정시간",
]

TEAM_OPTIONS = ["중계", "영상", "기타"]

SPORT_OPTIONS = [
    "프로야구",
    "KLPGA",
    "KPGA",
    "기타/ENG(실외)",
    "기타/ENG(실내)",
]

SPORT_COMMON_MEASURES = {
    "KPGA": (
        "- 중계차, 중계룸, 카메라룸, W/L룸, 몽골 텐트 냉방 가동\n"
        "  (냉방 공간 체감온도 27℃ 이하 유지)\n"
        "- 개인별 아이스박스·우산 지급 및 생수·얼음물 비치\n"
        "- 식염포도당·폭염질환 응급키트 위치 공유 및 사용 안내\n"
        "- 폭염 시간대 불필요한 외부 활동 최소화\n"
        "- 외부 근무자 1시간 이내 10분 이상 휴식 부여"
    ),
    "KLPGA": (
        "- 중계차, 중계룸, 카메라룸, W/L룸, 몽골 텐트 냉방 가동\n"
        "  (냉방 공간 체감온도 27℃ 이하 유지)\n"
        "- 개인별 아이스박스·우산 지급 및 생수·얼음물 비치\n"
        "- 식염포도당·폭염질환 응급키트 위치 공유 및 사용 안내\n"
        "- 폭염 시간대 불필요한 외부 활동 최소화\n"
        "- 외부 근무자 1시간 이내 10분 이상 휴식 부여"
    ),
    "프로야구": (
        "- 중계차·장비차·중계석·중계스태프실 냉방 가동\n"
        "  · 냉방 공간 체감온도 27℃ 이하 유지\n"
        "- 생수·냉음료·식염포도당·폭염질환 응급키트 비치 및 위치 공유\n"
        "- 폭염시간대 불필요한 야외활동 최소화\n"
        "- 이상 증상 발생 시 10~15분간 냉방 공간에서 휴식하도록\n"
        "  제작팀과 사전 협의"
    ),
    "기타/ENG(실내)": (
        "- 중계차, 휴게실, 체육관 냉방 가동\n"
        "  (냉방 공간 체감온도 27℃ 이하 유지)\n"
        "- 식염포도당, 폭염질환 응급키트 위치 공유 및 생수, 음료 지급\n"
        "- 폭염 시간대 불필요한 외부 활동 최소화\n"
        "- 중계차, 체육관 냉방 가동 공간에서 근무,\n"
        "  폭염 작업 해당 없음\n"
        "- 케이블 설치 등 외부 작업 완료\n"
        "  1시간 이내 10분 이상 휴게시간 부여"
    ),
    "기타/ENG(실외)": (
        "- 식염포도당, 폭염질환 응급키트 위치 공유 및 생수, 음료 지급\n"
        "- 폭염 시간대 불필요한 외부 활동 최소화\n"
        "- 중계차, 체육관 냉방 가동 공간에서 근무, 폭염 작업 해당 없음\n"
        "- 케이블 설치 등 외부 작업 완료 1시간 이내 10분 이상 휴게시간 부여"
    ),
}

# v3.16에서 조치사항에 함께 저장했던 기존 공통 문구.
# 마이그레이션 시 이 항목만 제거하고 실제 추가 조치는 보존합니다.
LEGACY_COMMON_MEASURES = {
    "야구": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "폭염시간대 업무강도 조정"
    ),
    "골프": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "코스 이동 동선 및 업무강도 조정"
    ),
    "남자골프": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "코스 이동 동선 및 업무강도 조정"
    ),
    "여자골프": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "코스 이동 동선 및 업무강도 조정"
    ),
    "기타스포츠": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "폭염시간대 업무강도 조정"
    ),
    "기타스포츠(실내)": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "폭염시간대 업무강도 조정"
    ),
    "기타스포츠(실외)": (
        "생수·냉수 비치 | 그늘·냉방 휴게공간 확인 | "
        "폭염시간대 업무강도 조정"
    ),
}

MEASURE_OPTIONS = [
    "1시간 이내 10분 이상 휴식",
    "2시간 이내 20분 이상 휴식",
    "냉방 장치 가동 공간에서 근무(29℃ 이하)",
    "근무 시간대 조정",
    "냉방 장치 가동 (중계차, 방송 시설 등)",
    "생수/식염포도당 등 제공",
    "휴식 공간(휴게실, 그늘막 등) 운영",
    "근무현장 의무시설 확인 및 공지",
]

TIME_ADJUST_REASON_PREFIX = "근무 시간대 조정 사유:"

TIME_FIELD_COLUMNS = {
    "work_start": "근무시작",
    "work_end": "근무종료",
    "heat_start": "폭염시작",
    "heat_end": "폭염종료",
}


WEATHER_LOGGER = logging.getLogger("checktemp.weather")
if not WEATHER_LOGGER.handlers:
    _weather_handler = logging.StreamHandler()
    _weather_handler.setFormatter(
        logging.Formatter(
            "%(asctime)s [CHECKTEMP-WEATHER] %(levelname)s %(message)s"
        )
    )
    WEATHER_LOGGER.addHandler(_weather_handler)
WEATHER_LOGGER.setLevel(logging.INFO)
WEATHER_LOGGER.propagate = False


def weather_debug_log_key(nonce: int) -> str:
    return f"weather_debug_log_{nonce}"


def clear_weather_debug_log(nonce: int) -> None:
    st.session_state[weather_debug_log_key(nonce)] = []


def add_weather_debug_log(
    nonce: int | None,
    message: str,
    *,
    level: str = "info",
) -> None:
    """기상조회 진단 로그를 Streamlit 로그와 화면용 세션에 함께 남깁니다."""
    now_text = datetime.now(KST).strftime("%H:%M:%S")
    clean_message = clean_text(message).replace("\n", " ")
    rendered = f"[{now_text}] {clean_message}"

    if level == "error":
        WEATHER_LOGGER.error(clean_message)
    elif level == "warning":
        WEATHER_LOGGER.warning(clean_message)
    else:
        WEATHER_LOGGER.info(clean_message)

    if nonce is None:
        return

    key = weather_debug_log_key(nonce)
    existing = list(st.session_state.get(key, []))
    existing.append(rendered)
    st.session_state[key] = existing[-80:]


def safe_error_body(error: urllib.error.HTTPError) -> str:
    """HTTP 오류 응답을 인증정보 없이 짧게 기록합니다."""
    try:
        payload = error.read(1200)
    except Exception:
        return ""

    if not payload:
        return ""

    for encoding in ("utf-8", "euc-kr"):
        try:
            return payload.decode(encoding, errors="replace")[:1000]
        except Exception:
            continue
    return repr(payload[:500])


st.set_page_config(
    page_title=APP_TITLE,
    page_icon=":material/health_and_safety:",
    layout="centered",
    initial_sidebar_state="collapsed",
)

st.markdown(
    """
    <style>
    :root {
        --bg: #f3f5f7;
        --surface: #ffffff;
        --surface-subtle: #f8f9fb;
        --ink: #111827;
        --muted: #667085;
        --line: #d8dee7;
        --line-strong: #c7ced8;
        --navy: #172b4d;
        --navy-hover: #0f213d;
        --danger: #b42318;
        --danger-soft: #fef3f2;
        --warning: #b54708;
        --warning-soft: #fff7ed;
        --success: #067647;
        --success-soft: #ecfdf3;
    }

    html, body, [class*="css"] {
        font-family: -apple-system, BlinkMacSystemFont, "Segoe UI",
                     "Noto Sans KR", sans-serif;
    }

    .stApp {
        background: var(--bg);
        color: var(--ink);
        color-scheme: light;
    }

    header[data-testid="stHeader"] {
        background: transparent;
    }

    #MainMenu, footer {
        visibility: hidden;
    }

    .block-container {
        max-width: 860px;
        padding-top: 1.35rem;
        padding-bottom: 5rem;
    }

    .app-header {
        display: flex;
        align-items: flex-start;
        justify-content: space-between;
        gap: 1.25rem;
        padding: 0.2rem 0 1.15rem;
        border-bottom: 1px solid var(--line);
        margin-bottom: 1rem;
    }

    .app-kicker {
        color: var(--muted);
        font-size: 0.74rem;
        font-weight: 700;
        letter-spacing: 0.08em;
        text-transform: uppercase;
        margin-bottom: 0.28rem;
    }

    .app-title {
        color: var(--ink);
        font-size: clamp(1.35rem, 3.2vw, 1.75rem) !important;
        font-weight: 760;
        letter-spacing: -0.035em;
        line-height: 1.2;
        margin: 0;
        white-space: nowrap;
    }

    .app-subtitle {
        color: var(--muted);
        font-size: 0.9rem;
        line-height: 1.55;
        margin: 0.42rem 0 0;
    }

    .app-version {
        flex: 0 0 auto;
        color: var(--muted);
        background: var(--surface-subtle);
        border: 1px solid var(--line);
        border-radius: 6px;
        padding: 0.38rem 0.55rem;
        font-size: 0.7rem;
        font-weight: 650;
        white-space: nowrap;
    }

    .section-heading {
        display: flex;
        align-items: flex-start;
        gap: 0.72rem;
        margin: 0 0 0.9rem;
    }

    .section-number {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        width: 1.75rem;
        height: 1.75rem;
        border-radius: 5px;
        background: var(--navy);
        color: #ffffff;
        font-size: 0.72rem;
        font-weight: 750;
        line-height: 1;
    }

    .section-heading b {
        display: block;
        color: var(--ink);
        font-size: 1rem;
        font-weight: 730;
        line-height: 1.3;
    }

    .section-heading small {
        display: block;
        color: var(--muted);
        font-size: 0.78rem;
        line-height: 1.45;
        margin-top: 0.12rem;
    }

    div[data-testid="stVerticalBlockBorderWrapper"] {
        background: var(--surface);
        border: 1px solid var(--line) !important;
        border-radius: 10px !important;
        box-shadow: 0 1px 2px rgba(16, 24, 40, 0.03);
    }

    div[data-testid="stVerticalBlockBorderWrapper"] > div {
        padding: 1.05rem 1.05rem 1.15rem;
    }

    div.stButton > button,
    div[data-testid="stDownloadButton"] > button,
    div[data-testid="stLinkButton"] a {
        min-height: 3.35rem;
        border-radius: 9px;
        border: 1px solid var(--line-strong);
        background: var(--surface);
        color: var(--ink);
        font-size: 1rem;
        font-weight: 720;
        box-shadow: none;
    }

    .field-label {
        color: #344054;
        font-size: 0.9rem;
        font-weight: 700;
        margin: 0.15rem 0 0.45rem;
    }

    div.stButton > button[kind="primary"] {
        background: var(--navy);
        border-color: var(--navy);
        color: #ffffff;
    }

    div.stButton > button[kind="primary"]:hover {
        background: var(--navy-hover);
        border-color: var(--navy-hover);
        color: #ffffff;
    }

    div[data-baseweb="input"] > div,
    div[data-baseweb="select"] > div,
    div[data-baseweb="textarea"] > div,
    div[data-testid="stTextInput"] input,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stTimeInput"] input,
    div[data-testid="stTextArea"] textarea,
    textarea {
        border-radius: 7px !important;
        border-color: var(--line-strong) !important;
        background-color: #ffffff !important;
        color: #111827 !important;
        -webkit-text-fill-color: #111827 !important;
    }

    div[data-baseweb="input"] input,
    div[data-baseweb="textarea"] textarea,
    div[data-baseweb="select"] *,
    div[data-testid="stTextInput"] input,
    div[data-testid="stNumberInput"] input,
    div[data-testid="stDateInput"] input,
    div[data-testid="stTimeInput"] input,
    div[data-testid="stTextArea"] textarea {
        color: #111827 !important;
        -webkit-text-fill-color: #111827 !important;
        caret-color: #111827 !important;
    }

    input::placeholder,
    textarea::placeholder {
        color: #98a2b3 !important;
        -webkit-text-fill-color: #98a2b3 !important;
        opacity: 1 !important;
    }

    div[data-baseweb="popover"],
    div[data-baseweb="popover"] > div,
    ul[role="listbox"],
    li[role="option"] {
        background-color: #ffffff !important;
        color: #111827 !important;
        -webkit-text-fill-color: #111827 !important;
    }

    li[role="option"] *,
    div[data-baseweb="menu"] *,
    div[data-baseweb="calendar"] * {
        color: #111827 !important;
        -webkit-text-fill-color: #111827 !important;
    }

    div[data-baseweb="calendar"],
    div[data-baseweb="calendar"] > div,
    div[data-baseweb="calendar"] button {
        background-color: #ffffff !important;
    }

    div[data-baseweb="calendar"] button[aria-selected="true"] {
        background-color: var(--navy) !important;
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
    }

    div[data-baseweb="calendar"] button[aria-selected="true"] * {
        color: #ffffff !important;
        -webkit-text-fill-color: #ffffff !important;
    }

    label[data-testid="stWidgetLabel"] p {
        color: #344054;
        font-size: 0.84rem;
        font-weight: 650;
    }

    .quick-time-summary,
    .metric-grid,
    .record-details {
        display: grid;
        grid-template-columns: repeat(3, minmax(0, 1fr));
        gap: 0.62rem;
    }

    .quick-time-summary {
        margin: 0.72rem 0 0.7rem;
    }

    .metric-grid {
        margin: 0.65rem 0 1rem;
    }

    .record-details {
        gap: 0.48rem;
        margin: 0.8rem 0 0.75rem;
    }

    .quick-time-card,
    .metric-card,
    .record-detail {
        background: var(--surface-subtle);
        border: 1px solid var(--line);
        border-radius: 8px;
        padding: 0.78rem 0.82rem;
    }

    .metric-card {
        background: var(--surface);
        padding: 0.9rem 0.95rem;
    }

    .quick-time-card span,
    .metric-card span,
    .record-detail span {
        display: block;
        color: var(--muted);
        font-size: 0.72rem;
        font-weight: 650;
        margin-bottom: 0.2rem;
    }

    .quick-time-card strong {
        display: block;
        color: var(--ink);
        font-size: 0.97rem;
        font-weight: 720;
        white-space: nowrap;
    }

    .metric-card b {
        display: block;
        color: var(--ink);
        font-size: 1.45rem;
        font-weight: 740;
        line-height: 1.15;
    }

    .record-detail b {
        display: block;
        color: var(--ink);
        font-size: 0.84rem;
        font-weight: 680;
        line-height: 1.35;
    }

    .status-pill {
        display: inline-flex;
        align-items: center;
        min-height: 1.7rem;
        padding: 0.18rem 0.5rem;
        border-radius: 5px;
        border: 1px solid var(--line);
        background: var(--surface-subtle);
        color: #475467;
        font-size: 0.72rem;
        font-weight: 720;
        margin-bottom: 0.42rem;
    }

    .status-pill.status-caution {
        background: var(--warning-soft);
        border-color: #fed7aa;
        color: var(--warning);
    }

    .status-pill.status-danger {
        background: var(--danger-soft);
        border-color: #fecdca;
        color: var(--danger);
    }

    .status-pill.status-normal {
        background: var(--success-soft);
        border-color: #abefc6;
        color: var(--success);
    }

    .setup-box {
        background: #fffcf5;
        border: 1px solid #fedf89;
        border-radius: 8px;
        color: #7a2e0e;
        padding: 0.9rem 1rem;
        margin-bottom: 1rem;
    }

    @media (max-width: 600px) {
        .block-container {
            padding-left: 0.78rem;
            padding-right: 0.78rem;
            padding-top: 0.85rem;
        }

        .app-header {
            display: block;
        }

        .app-title {
            font-size: clamp(0.78rem, 4.2vw, 1.15rem) !important;
            letter-spacing: -0.07em;
            width: calc(100vw - 1.1rem);
            max-width: calc(100vw - 1.1rem);
            white-space: nowrap;
        }

        .app-version {
            display: inline-block;
            margin-top: 0.7rem;
        }

        .quick-time-summary,
        .metric-grid,
        .record-details {
            grid-template-columns: 1fr;
        }

        div[data-testid="stVerticalBlockBorderWrapper"] > div {
            padding: 0.9rem;
        }

        div.stButton > button,
        div[data-testid="stDownloadButton"] > button,
        div[data-testid="stLinkButton"] a {
            min-height: 3.6rem;
            font-size: 1.03rem;
        }
    }
    </style>
    """,
    unsafe_allow_html=True,
)


def empty_dataframe() -> pd.DataFrame:
    return pd.DataFrame(columns=COLUMNS)


def clean_text(value: Any) -> str:
    if value is None:
        return ""
    if isinstance(value, float) and pd.isna(value):
        return ""
    return str(value).strip()


@st.cache_data(show_spinner=False, ttl=86400)
def search_kakao_site_identity(
    site_name: str,
    rest_api_key: str,
) -> dict[str, Any]:
    """카카오 장소검색의 장소 ID·표준명·좌표를 반환합니다."""
    params = urllib.parse.urlencode({"query": site_name, "size": "5"})
    request = urllib.request.Request(
        f"{KAKAO_PLACE_API_URL}?{params}",
        headers={
            "Authorization": f"KakaoAK {rest_api_key}",
            "User-Agent": "checktemp-streamlit/1.0",
        },
    )
    with urllib.request.urlopen(request, timeout=8) as response:
        result = json.loads(response.read().decode("utf-8"))
    documents = result.get("documents") or []
    if not documents:
        raise ValueError("카카오 장소 검색 결과가 없습니다")
    place = documents[0]
    return {
        "id": clean_text(place.get("id")),
        "name": clean_text(place.get("place_name")) or site_name,
        "address": (
            clean_text(place.get("road_address_name"))
            or clean_text(place.get("address_name"))
        ),
        "latitude": float(place["y"]),
        "longitude": float(place["x"]),
    }


def generic_report_site_identity(site_name: Any) -> tuple[str, str]:
    """새 장소도 표기 차이만으로 보고서가 나뉘지 않도록 비교 키를 만듭니다."""
    original = clean_text(site_name)
    display = re.sub(r"\s+", " ", original).strip()
    display = re.sub(r"\s*[\\(（][^\\)）]*[\\)）]\s*", " ", display).strip()
    display = display.replace("센타", "센터")

    folded = display.casefold()
    folded = folded.replace("씨씨", "cc")
    compact = re.sub(r"[\s\-_/.,·ㆍ:：]+", "", folded)
    compact = compact.replace("컨트리클럽", "cc")
    compact = compact.replace("골프클럽", "gc")
    compact = compact.replace("골프장", "gc")
    compact = compact.replace("센타", "센터")

    key = re.sub(r"(?:cc|gc)$", "", compact)
    key = key or compact or original.casefold()

    golf_suffix = re.search(
        r"(?:cc|씨씨|컨트리클럽|골프클럽|골프장)\s*$",
        display,
        flags=re.IGNORECASE,
    )
    if golf_suffix:
        display = re.sub(
            r"\s*(?:cc|씨씨|컨트리클럽|골프클럽|골프장)\s*$",
            "",
            display,
            flags=re.IGNORECASE,
        ).strip()
        display = f"{display}CC"

    return f"normalized:{key}", display or original


def report_place_api_key() -> str:
    """보고서 장소 묶기에 사용할 카카오 REST API 키를 찾습니다."""
    for env_name in ("KAKAO_REST_API_KEY", "KAKAO_REST_KEY"):
        value = clean_text(os.environ.get(env_name))
        if value:
            return value
    return clean_text(get_secret(("location", "kakao_rest_api_key"), ""))


REPORT_SITE_ALIAS_RULES = (
    ("몽베르CC", ("몽베르",), ()),
    ("솔라고CC", ("솔라고",), ()),
    (
        "포천힐스CC",
        ("포천힐스", "포천힐스cc", "포천힐스씨씨", "포천힐스컨트리클럽"),
        (),
    ),
    ("사직야구장", ("사직",), ("야구장", "구장")),
    (
        "단양군체육관",
        ("단양",),
        ("군체육관", "국민체육센타", "국민체육센터", "체육관"),
    ),
    ("광주기아챔피언스필드", ("챔피언스필드",), ("광주", "기아")),
    ("창원NC파크", ("창원nc파크", "창원엔씨파크"), ()),
)


def report_site_alias_key(value: Any) -> str:
    """장소 별칭 비교용 문자열을 만듭니다."""
    text = clean_text(value).casefold()
    text = text.replace("센타", "센터").replace("씨씨", "cc")
    text = re.sub(r"\s*[\\(（][^\\)）]*[\\)）]\s*", "", text)
    text = re.sub(r"[\s\-_/.,·ㆍ:：]+", "", text)
    text = text.replace("컨트리클럽", "cc")
    text = text.replace("골프클럽", "gc")
    text = text.replace("골프장", "gc")
    return text


def known_report_site_identity(site_name: Any) -> tuple[str, str] | None:
    """API보다 우선 적용할 현장 별칭을 찾습니다."""
    alias_key = report_site_alias_key(site_name)
    if not alias_key:
        return None
    for display_name, aliases, required_any in REPORT_SITE_ALIAS_RULES:
        alias_matches = any(
            report_site_alias_key(alias) in alias_key
            for alias in aliases
        )
        required_matches = (
            not required_any
            or any(report_site_alias_key(token) in alias_key for token in required_any)
        )
        if alias_matches and required_matches:
            keyword = report_site_alias_key(display_name)
            return f"keyword:{keyword}", display_name
    return None


def report_site_identity(
    site_name: Any,
    kakao_key: str = "",
) -> tuple[str, str]:
    """장소 API, 지정 키워드, 일반 정규화로 보고서 장소를 구분합니다."""
    original = clean_text(site_name)
    known_identity = known_report_site_identity(original)
    if known_identity:
        return known_identity

    if kakao_key and original:
        try:
            place = search_kakao_site_identity(original, kakao_key)
            place_id = clean_text(place.get("id"))
            place_name = clean_text(place.get("name")) or original
            if place_id:
                display_name = generic_report_site_identity(place_name)[1]
                return f"kakao-place:{place_id}", display_name
        except Exception:  # noqa: BLE001
            pass

    return generic_report_site_identity(original)


def report_site_similarity_key(value: Any) -> str:
    """유사 장소 판단에 사용할 느슨한 비교 문자열을 만듭니다."""
    text = clean_text(value).casefold()
    text = text.replace("센타", "센터").replace("씨씨", "cc")
    text = re.sub(r"\s*[\\(（][^\\)）]*[\\)）]\s*", "", text)
    text = re.sub(r"[\s\-_/.,·ㆍ:：]+", "", text)
    for suffix in (
        "컨트리클럽",
        "골프클럽",
        "골프장",
        "국민체육센터",
        "문화체육센터",
        "실내체육관",
        "체육센터",
        "체육관",
        "주차장",
        "센터",
        "cc",
        "gc",
    ):
        text = text.replace(suffix, "")
    return text


def report_site_name_similarity(left: Any, right: Any) -> float:
    """두 장소명이 같은 현장일 가능성을 0~1로 계산합니다."""
    left_key = report_site_similarity_key(left)
    right_key = report_site_similarity_key(right)
    if not left_key or not right_key:
        return 0.0
    if left_key == right_key:
        return 1.0
    sequence_score = difflib.SequenceMatcher(None, left_key, right_key).ratio()
    left_chars = set(left_key)
    right_chars = set(right_key)
    dice_score = (
        2 * len(left_chars & right_chars) / (len(left_chars) + len(right_chars))
        if left_chars and right_chars
        else 0.0
    )
    return max(sequence_score, dice_score)


def report_site_region_key(value: Any) -> str:
    """종목과 함께 묶을 지역 단위 키를 만듭니다."""
    text = report_site_alias_key(value)
    if not text:
        return ""
    explicit_region = re.search(r"([가-힣]{2,4})(?:시|군|구)", text)
    if explicit_region:
        return explicit_region.group(1)
    hangul_parts = re.findall(r"[가-힣]+", text)
    if not hangul_parts:
        return ""
    first_part = hangul_parts[0]
    if len(first_part) >= 2:
        return first_part[:2]
    return first_part


def report_rows_sport_region_key(rows: pd.DataFrame) -> str:
    """보고서 행 묶음의 종목+지역 키를 반환합니다."""
    if rows.empty:
        return ""
    first = rows.iloc[0]
    site = clean_text(first.get("현장명"))
    sport = normalize_sport(first.get("종목"), site)
    if sport not in {"프로야구", "KLPGA", "KPGA"}:
        return ""
    region = report_site_region_key(site)
    if not region:
        return ""
    return f"{sport}:{region}"


def report_sport_region_match(
    left_rows: pd.DataFrame,
    right_rows: pd.DataFrame,
) -> bool:
    """업무내용과 지역이 같으면 표기 차이로 보고서가 갈라지지 않게 합니다."""
    left_key = report_rows_sport_region_key(left_rows)
    right_key = report_rows_sport_region_key(right_rows)
    return bool(left_key and left_key == right_key)


def report_time_span_minutes(rows: pd.DataFrame) -> tuple[int | None, int | None]:
    """기록 묶음의 가장 이른 시작과 가장 늦은 종료를 분 단위로 반환합니다."""
    starts: list[int] = []
    ends: list[int] = []
    for value in rows.get("근무시작", pd.Series(dtype=object)).tolist():
        parsed = parse_time_value(value)
        if parsed is not None:
            starts.append(parsed.hour * 60 + parsed.minute)
    for value in rows.get("근무종료", pd.Series(dtype=object)).tolist():
        parsed = parse_time_value(value)
        if parsed is not None:
            ends.append(parsed.hour * 60 + parsed.minute)
    return (min(starts) if starts else None, max(ends) if ends else None)


def report_time_spans_close(
    left_rows: pd.DataFrame,
    right_rows: pd.DataFrame,
) -> bool:
    """근무시간이 겹치거나 시작 시간이 크게 다르지 않은지 확인합니다."""
    left_start, left_end = report_time_span_minutes(left_rows)
    right_start, right_end = report_time_span_minutes(right_rows)
    if None in (left_start, left_end, right_start, right_end):
        return True
    assert left_start is not None
    assert left_end is not None
    assert right_start is not None
    assert right_end is not None
    overlaps = max(left_start, right_start) <= min(left_end, right_end)
    starts_close = abs(left_start - right_start) <= 180
    return overlaps or starts_close


def report_broadcast_team_pair(
    left_rows: pd.DataFrame,
    right_rows: pd.DataFrame,
) -> bool:
    """중계·영상이 함께 움직인 기록인지 확인합니다."""
    left_teams = {
        clean_text(value).removesuffix("팀")
        for value in left_rows.get("팀", pd.Series(dtype=object)).tolist()
    }
    right_teams = {
        clean_text(value).removesuffix("팀")
        for value in right_rows.get("팀", pd.Series(dtype=object)).tolist()
    }
    combined = left_teams | right_teams
    return "중계" in combined and "영상" in combined


def merge_similar_report_sites(normalized: pd.DataFrame) -> pd.DataFrame:
    """같은 날짜의 중계·영상 유사 장소를 보고서 한 묶음으로 합칩니다."""
    if normalized.empty:
        return normalized

    for work_date, date_rows in normalized.groupby("작업날짜", dropna=False):
        _ = work_date
        groups = [
            (key, rows.copy())
            for key, rows in date_rows.groupby("_보고서장소키", dropna=False)
        ]
        if len(groups) < 2:
            continue

        parent = {key: key for key, _ in groups}

        def find(key: str) -> str:
            while parent[key] != key:
                parent[key] = parent[parent[key]]
                key = parent[key]
            return key

        def union(left: str, right: str) -> None:
            left_root = find(left)
            right_root = find(right)
            if left_root != right_root:
                parent[right_root] = left_root

        for left_index, (left_key, left_rows) in enumerate(groups):
            for right_key, right_rows in groups[left_index + 1:]:
                if not report_broadcast_team_pair(left_rows, right_rows):
                    continue
                if not report_time_spans_close(left_rows, right_rows):
                    continue
                if report_sport_region_match(left_rows, right_rows):
                    union(left_key, right_key)
                    continue
                left_name = clean_text(left_rows.iloc[0].get("현장명"))
                right_name = clean_text(right_rows.iloc[0].get("현장명"))
                similarity = report_site_name_similarity(left_name, right_name)
                if similarity >= 0.58:
                    union(left_key, right_key)

        merged_names: dict[str, str] = {}
        for key, rows in groups:
            root = find(key)
            names = unique_texts(rows["현장명"].tolist())
            current = merged_names.get(root, "")
            for name in names:
                if not current:
                    current = name
                    continue
                current_has_cc = current.casefold().endswith("cc")
                candidate_has_cc = name.casefold().endswith("cc")
                if candidate_has_cc and not current_has_cc:
                    current = name
                elif candidate_has_cc == current_has_cc and len(name) < len(current):
                    current = name
            merged_names[root] = current

        date_index = date_rows.index
        normalized.loc[date_index, "_보고서장소키"] = normalized.loc[
            date_index,
            "_보고서장소키",
        ].map(lambda key: f"similar:{find(key)}")
        normalized.loc[date_index, "현장명"] = normalized.loc[
            date_index,
            "_보고서장소키",
        ].map(lambda key: merged_names.get(key.removeprefix("similar:"), ""))

    return normalized


def canonicalize_report_rows(
    records: pd.DataFrame,
) -> pd.DataFrame:
    """원본 기록을 바꾸지 않고 보고서용 장소 키·명칭을 적용합니다."""
    normalized = records.copy()
    kakao_key = report_place_api_key()
    site_names = list(dict.fromkeys(
        clean_text(site) for site in normalized["현장명"].tolist()
    ))
    identities = {
        site: report_site_identity(site, kakao_key)
        for site in site_names
    }
    preferred_names: dict[str, str] = {}
    for key, display_name in identities.values():
        current = preferred_names.get(key, "")
        if not current:
            preferred_names[key] = display_name
            continue
        current_has_cc = current.casefold().endswith("cc")
        candidate_has_cc = display_name.casefold().endswith("cc")
        if candidate_has_cc and not current_has_cc:
            preferred_names[key] = display_name
        elif candidate_has_cc == current_has_cc and len(display_name) < len(current):
            preferred_names[key] = display_name
    normalized["_보고서장소키"] = normalized["현장명"].map(
        lambda site: identities[clean_text(site)][0]
    )
    normalized["현장명"] = normalized["현장명"].map(
        lambda site: preferred_names[identities[clean_text(site)][0]]
    )
    normalized = merge_similar_report_sites(normalized)
    return normalized


def parse_int(value: Any, default: int = 0) -> int:
    try:
        return int(float(clean_text(value)))
    except (TypeError, ValueError):
        return default


def parse_float(value: Any) -> float | None:
    try:
        text = clean_text(value)
        return float(text) if text else None
    except (TypeError, ValueError):
        return None


def format_number(value: float) -> str:
    if value.is_integer():
        return str(int(value))
    return f"{value:.1f}".rstrip("0").rstrip(".")


def normalize_temperature_text(value: Any) -> tuple[str, str | None]:
    """
    체감온도 입력값을 검증하고 저장용 문자열로 정리합니다.

    허용 예:
    - 34
    - 33~35
    - 33.5~35.2
    - 33-35
    - 33～35℃
    """
    text = clean_text(value)
    if not text:
        return "", None

    text = (
        text.replace("°C", "")
        .replace("°c", "")
        .replace("℃", "")
        .replace("도", "")
        .replace("～", "~")
        .replace("〜", "~")
        .replace("–", "~")
        .replace("—", "~")
    )

    # 양수 두 값 사이의 일반 하이픈은 범위 기호로 처리합니다.
    text = re.sub(r"(?<=\d)\s*-\s*(?=\d)", "~", text)
    text = re.sub(r"\s+", "", text)

    match = re.fullmatch(
        r"(-?\d+(?:\.\d+)?)(?:~(-?\d+(?:\.\d+)?))?",
        text,
    )
    if not match:
        return "", "체감온도는 34 또는 33~35 형식으로 입력해 주세요."

    first = float(match.group(1))
    second_text = match.group(2)
    second = float(second_text) if second_text is not None else None

    values = [first] if second is None else [first, second]
    if any(number < -20 or number > 60 for number in values):
        return "", "체감온도는 -20℃에서 60℃ 사이로 입력해 주세요."

    if second is not None and first > second:
        return "", "체감온도 범위는 낮은 값부터 입력해 주세요. 예: 33~35"

    normalized = format_number(first)
    if second is not None:
        normalized += f"~{format_number(second)}"

    return normalized, None


def temperature_numbers(value: Any) -> list[float]:
    normalized, error = normalize_temperature_text(value)
    if error or not normalized:
        return []

    numbers: list[float] = []
    for part in normalized.split("~"):
        try:
            numbers.append(float(part))
        except ValueError:
            continue
    return numbers


def max_temperature(value: Any) -> float | None:
    numbers = temperature_numbers(value)
    return max(numbers) if numbers else None


def temperature_display(value: Any) -> str:
    normalized, error = normalize_temperature_text(value)
    if error:
        return clean_text(value)
    return normalized


def parse_time_value(value: Any, default: time | None = None) -> time | None:
    """시간을 HH:MM 형식으로 해석합니다. 930→09:30, 18→18:00도 허용합니다."""
    text = clean_text(value)
    if not text:
        return default

    normalized = text.replace("시", ":00").replace(" ", "")

    if re.fullmatch(r"\d{1,2}", normalized):
        normalized = f"{normalized}:00"
    elif re.fullmatch(r"\d{3}", normalized):
        normalized = f"0{normalized[0]}:{normalized[1:]}"
    elif re.fullmatch(r"\d{4}", normalized):
        normalized = f"{normalized[:2]}:{normalized[2:]}"

    for fmt in ("%H:%M", "%H:%M:%S"):
        try:
            return datetime.strptime(normalized, fmt).time().replace(
                second=0,
                microsecond=0,
            )
        except ValueError:
            continue
    return default


def format_time_value(value: time | None) -> str:
    return value.strftime("%H:%M") if value is not None else ""


def markdown_escape(value: Any) -> str:
    return re.sub(
        r"([\\`*_{}\[\]()#+\-.!|>])",
        r"\\\1",
        clean_text(value),
    )


def get_secret(path: tuple[str, ...], default: Any = "") -> Any:
    try:
        current: Any = st.secrets
        for key in path:
            current = current[key]
        if hasattr(current, "items") or isinstance(current, (list, tuple)):
            return current
        return clean_text(current)
    except (KeyError, TypeError):
        return default
    except Exception as exc:
        if exc.__class__.__name__ == "StreamlitSecretNotFoundError":
            return default
        raise


def time_state_key(field: str, nonce: int) -> str:
    return f"time_{field}_{nonce}"


def initialize_time_state(
    editing_record: dict[str, Any],
    nonce: int,
) -> None:
    for field, column in TIME_FIELD_COLUMNS.items():
        key = time_state_key(field, nonce)
        if key not in st.session_state:
            st.session_state[key] = parse_time_value(
                editing_record.get(column)
            )


def set_time_now(field: str, nonce: int) -> None:
    st.session_state[time_state_key(field, nonce)] = (
        datetime.now(KST).time().replace(second=0, microsecond=0)
    )


def weather_notice_key(nonce: int) -> str:
    return f"weather_notice_{nonce}"


@st.cache_resource(show_spinner=False)
def persistent_weather_results() -> dict[str, dict[str, Any]]:
    """앱 재접속 시에도 같은 조회 조건의 결과를 재사용합니다."""
    return {}


def weather_result_cache_key(nonce: int) -> str:
    selected_date = st.session_state.get(f"date_{nonce}")
    date_text = (
        selected_date.isoformat()
        if isinstance(selected_date, date)
        else clean_text(selected_date)
    )
    values = [
        date_text,
        clean_text(st.session_state.get(f"site_{nonce}")).casefold(),
        clean_text(st.session_state.get(f"manual_work_start_{nonce}")),
        clean_text(st.session_state.get(f"manual_work_end_{nonce}")),
        clean_text(st.session_state.get(f"manual_heat_start_{nonce}")),
        clean_text(st.session_state.get(f"manual_heat_end_{nonce}")),
    ]
    return "|".join(values)


def save_weather_result_cache(nonce: int) -> None:
    notice = st.session_state.get(weather_notice_key(nonce))
    temperature = clean_text(st.session_state.get(f"temperature_{nonce}"))
    cache_key = weather_result_cache_key(nonce)
    if not cache_key.strip("|") or not temperature or not notice:
        return
    cache = persistent_weather_results()
    cache[cache_key] = {
        "temperature": temperature,
        "notice": notice,
        "saved_at": datetime.now(KST).timestamp(),
    }
    # 서버 메모리가 불필요하게 커지지 않도록 최근 300건만 유지합니다.
    if len(cache) > 300:
        oldest_key = min(cache, key=lambda key: cache[key]["saved_at"])
        cache.pop(oldest_key, None)


def restore_weather_result_cache(nonce: int) -> bool:
    cache_key = weather_result_cache_key(nonce)
    cached = persistent_weather_results().get(cache_key)
    if not cached:
        return False
    # 12시간 동안만 재사용하여 오래된 관측값이 남지 않게 합니다.
    if datetime.now(KST).timestamp() - float(cached["saved_at"]) > 43200:
        persistent_weather_results().pop(cache_key, None)
        return False
    if weather_notice_key(nonce) not in st.session_state:
        st.session_state[weather_notice_key(nonce)] = cached["notice"]
    if f"temperature_{nonce}" not in st.session_state:
        st.session_state[f"temperature_{nonce}"] = cached["temperature"]
    return True


def get_site_coordinates(site_name: str) -> tuple[float, float] | None:
    """Streamlit Secrets에서 현장명의 위도·경도를 찾습니다."""
    sites = get_secret(("weather", "sites"), default=None)
    if not hasattr(sites, "items"):
        return None

    target = clean_text(site_name).casefold()

    for configured_name, configured_value in sites.items():
        if clean_text(configured_name).casefold() != target:
            continue

        try:
            if isinstance(configured_value, str):
                parts = [
                    part.strip()
                    for part in configured_value.split(",")
                ]
                latitude, longitude = map(float, parts[:2])
            elif hasattr(configured_value, "get"):
                latitude = float(configured_value.get("lat"))
                longitude = float(configured_value.get("lon"))
            else:
                latitude = float(configured_value[0])
                longitude = float(configured_value[1])
        except (IndexError, TypeError, ValueError):
            return None

        if not (33 <= latitude <= 39 and 124 <= longitude <= 132):
            return None

        return latitude, longitude

    return None


def search_kakao_site_coordinates(
    site_name: str,
    rest_api_key: str,
) -> tuple[float, float, str]:
    """카카오 장소 검색으로 현장명의 위도·경도를 찾습니다."""
    place = search_kakao_site_identity(site_name, rest_api_key)
    latitude = float(place["latitude"])
    longitude = float(place["longitude"])
    matched_name = clean_text(place.get("name")) or site_name

    if not (33 <= latitude <= 39 and 124 <= longitude <= 132):
        raise ValueError("검색된 장소가 대한민국 범위를 벗어났습니다")

    return latitude, longitude, matched_name


def parse_kma_temperature_observations(
    response_text: str,
) -> list[tuple[str, float]]:
    """기상청 특정지점 ASCII 응답에서 체감온도 관측값을 추출합니다."""
    observations: list[tuple[str, float]] = []

    for raw_line in response_text.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("#"):
            continue

        # 응답 버전에 따라 공백/쉼표 구분 및 초(14자리) 포함 여부가
        # 달라질 수 있으므로 행 전체에서 시각과 값을 찾습니다.
        timestamp_match = re.search(r"(?<!\d)(\d{12}|\d{14})(?!\d)", line)
        if timestamp_match:
            timestamp = timestamp_match.group(1)[:12]
            value_text = line[timestamp_match.end():]
        else:
            separated_time = re.search(
                r"(\d{4})[-/.]?(\d{2})[-/.]?(\d{2})"
                r"[ T]?(\d{2}):?(\d{2})",
                line,
            )
            if not separated_time:
                continue
            timestamp = "".join(separated_time.groups())
            value_text = line[separated_time.end():]

        numeric_values = re.findall(r"[-+]?\d+(?:\.\d+)?", value_text)
        if not numeric_values:
            continue

        # 특정지점 단일요소 응답의 관측값은 행의 마지막 숫자입니다.
        value = float(numeric_values[-1])
        if -50 <= value <= 80 and value not in (-99.0, -999.0, -9999.0):
            observations.append((timestamp, value))

    if not observations:
        raise ValueError("기상청 응답에서 체감온도 값을 찾지 못했습니다.")

    observations.sort(key=lambda item: item[0])
    return observations


def parse_kma_temperature_response(response_text: str) -> tuple[str, str]:
    """기상청 특정지점 ASCII 응답에서 최신 체감온도를 추출합니다."""
    observations = parse_kma_temperature_observations(response_text)

    timestamp, value = observations[-1]
    observed_at = datetime.strptime(timestamp, "%Y%m%d%H%M").strftime(
        "%H:%M"
    )
    return format_number(value), observed_at


def fetch_kma_apparent_temperature(
    latitude: float,
    longitude: float,
    auth_key: str,
) -> tuple[str, str]:
    """기상청 500m 격자 특정지점의 최신 체감온도를 조회합니다."""
    current_time = datetime.now(KST).replace(second=0, microsecond=0)
    query_end = current_time - timedelta(
        minutes=(current_time.minute % 5) + 10
    )
    query_start = query_end - timedelta(minutes=55)
    params = urllib.parse.urlencode(
        {
            "obs": "ta_chi",
            "tm1": query_start.strftime("%Y%m%d%H%M"),
            "tm2": query_end.strftime("%Y%m%d%H%M"),
            "itv": "5",
            "lon": f"{longitude:.6f}",
            "lat": f"{latitude:.6f}",
            "authKey": auth_key,
        }
    )
    request = urllib.request.Request(
        f"{KMA_POINT_API_URL}?{params}",
        headers={"User-Agent": "checktemp-streamlit/1.0"},
    )

    with urllib.request.urlopen(request, timeout=8) as response:
        payload = response.read()

    for encoding in ("utf-8", "euc-kr"):
        try:
            response_text = payload.decode(encoding)
            break
        except UnicodeDecodeError:
            continue
    else:
        response_text = payload.decode("utf-8", errors="replace")

    return parse_kma_temperature_response(response_text)


def fetch_kma_apparent_temperature_range(
    latitude: float,
    longitude: float,
    auth_key: str,
    range_start: datetime,
    range_end: datetime,
    *,
    debug_nonce: int | None = None,
) -> dict[str, Any]:
    """500m 격자 체감온도를 55분 이하 구간으로 독립 조회합니다.

    일부 구간이 실패해도 성공한 관측값은 보존합니다.
    연속 네트워크 타임아웃이 여러 차례 발생하면 500m 조회를 중단하고
    초단기실황 조회로 넘어갈 수 있도록 부분 결과를 반환합니다.
    """
    now = datetime.now(KST).replace(second=0, microsecond=0)

    if range_start > now:
        raise ValueError("조회 시작시간이 현재보다 이후입니다.")

    actual_end = min(range_end, now)
    if actual_end < range_start:
        raise ValueError("조회할 시간대가 없습니다.")

    add_weather_debug_log(
        debug_nonce,
        (
            "500m 분할조회 시작 | "
            f"range={range_start.strftime('%Y-%m-%d %H:%M')}"
            f"~{actual_end.strftime('%Y-%m-%d %H:%M')} | "
            f"lat={latitude:.6f}, lon={longitude:.6f}"
        ),
    )

    observations_by_time: dict[str, float] = {}
    failed_chunks: list[str] = []
    requested_chunks = 0
    successful_chunks = 0
    consecutive_timeouts = 0

    chunk_start = range_start

    while chunk_start <= actual_end:
        # 공식 최대 조회기간보다 여유를 둔 55분 단위.
        chunk_end = min(
            chunk_start + timedelta(minutes=55),
            actual_end,
        )
        requested_chunks += 1
        chunk_label = (
            f"{chunk_start.strftime('%H:%M')}~"
            f"{chunk_end.strftime('%H:%M')}"
        )

        params = urllib.parse.urlencode(
            {
                "obs": "ta_chi",
                "tm1": chunk_start.strftime("%Y%m%d%H%M"),
                "tm2": chunk_end.strftime("%Y%m%d%H%M"),
                "itv": "5",
                "lon": f"{longitude:.6f}",
                "lat": f"{latitude:.6f}",
                "authKey": auth_key,
            }
        )

        request = urllib.request.Request(
            f"{KMA_POINT_API_URL}?{params}",
            headers={"User-Agent": "checktemp-streamlit/1.0"},
        )

        add_weather_debug_log(
            debug_nonce,
            (
                f"500m 구간 요청 {requested_chunks} | {chunk_label} | "
                f"timeout={KMA_POINT_RANGE_TIMEOUT_SECONDS}s"
            ),
        )

        started = time_module.perf_counter()
        payload: bytes | None = None
        network_timeout = False

        try:
            with urllib.request.urlopen(
                request,
                timeout=KMA_POINT_RANGE_TIMEOUT_SECONDS,
            ) as response:
                payload = response.read()
                elapsed = time_module.perf_counter() - started
                status = getattr(response, "status", None) or response.getcode()

            add_weather_debug_log(
                debug_nonce,
                (
                    f"500m 구간 성공 | {chunk_label} | status={status} | "
                    f"elapsed={elapsed:.2f}s | bytes={len(payload)}"
                ),
            )
            consecutive_timeouts = 0

        except urllib.error.HTTPError as error:
            elapsed = time_module.perf_counter() - started
            body = safe_error_body(error)
            failed_chunks.append(chunk_label)
            consecutive_timeouts = 0
            add_weather_debug_log(
                debug_nonce,
                (
                    f"500m 구간 HTTPError | {chunk_label} | "
                    f"code={error.code} | elapsed={elapsed:.2f}s"
                    + (f" | body={body}" if body else "")
                ),
                level="error",
            )

        except (urllib.error.URLError, TimeoutError) as error:
            elapsed = time_module.perf_counter() - started
            failed_chunks.append(chunk_label)
            consecutive_timeouts += 1
            network_timeout = True
            reason = clean_text(getattr(error, "reason", error))
            add_weather_debug_log(
                debug_nonce,
                (
                    f"500m 구간 타임아웃 | {chunk_label} | "
                    f"type={type(error).__name__} | reason={reason} | "
                    f"elapsed={elapsed:.2f}s | "
                    f"consecutive={consecutive_timeouts}"
                ),
                level="error",
            )

        except Exception as error:
            elapsed = time_module.perf_counter() - started
            failed_chunks.append(chunk_label)
            consecutive_timeouts = 0
            add_weather_debug_log(
                debug_nonce,
                (
                    f"500m 구간 예외 | {chunk_label} | "
                    f"type={type(error).__name__} | "
                    f"elapsed={elapsed:.2f}s | message={error}"
                ),
                level="error",
            )

        if payload is not None:
            response_text = ""
            for encoding in ("utf-8", "euc-kr"):
                try:
                    response_text = payload.decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue

            if not response_text:
                response_text = payload.decode(
                    "utf-8",
                    errors="replace",
                )

            try:
                chunk_observations = parse_kma_temperature_observations(
                    response_text
                )
                for timestamp, value in chunk_observations:
                    observations_by_time[timestamp] = value

                successful_chunks += 1
                add_weather_debug_log(
                    debug_nonce,
                    (
                        f"500m 구간 파싱 성공 | {chunk_label} | "
                        f"observations={len(chunk_observations)}"
                    ),
                )
            except Exception as error:
                failed_chunks.append(chunk_label)
                preview = re.sub(r"\s+", " ", response_text[:600])
                add_weather_debug_log(
                    debug_nonce,
                    (
                        f"500m 구간 파싱 실패 | {chunk_label} | "
                        f"type={type(error).__name__} | "
                        f"message={error} | preview={preview}"
                    ),
                    level="error",
                )

        # 서버가 연속으로 응답하지 않으면 이후 구간까지 오래 기다리지 않습니다.
        if (
            network_timeout
            and consecutive_timeouts >= KMA_POINT_RANGE_MAX_CONSECUTIVE_TIMEOUTS
        ):
            add_weather_debug_log(
                debug_nonce,
                (
                    "500m 회로차단 | "
                    f"연속 {KMA_POINT_RANGE_MAX_CONSECUTIVE_TIMEOUTS}개 "
                    "구간 타임아웃으로 "
                    "나머지 500m 조회를 중단하고 지역실황으로 진행"
                ),
                level="warning",
            )
            break

        if chunk_end >= actual_end:
            break

        # itv=5에 맞춰 다음 관측 시점부터 시작.
        chunk_start = chunk_end + timedelta(minutes=5)

    values = list(observations_by_time.values())

    if not values:
        add_weather_debug_log(
            debug_nonce,
            (
                f"500m 분할조회 결과 없음 | requested={requested_chunks} | "
                f"failed={len(failed_chunks)}"
            ),
            level="warning",
        )
        return {
            "available": False,
            "source": "500m 격자",
            "minimum": None,
            "maximum": None,
            "actual_end": actual_end.strftime("%H:%M"),
            "observations": 0,
            "requested": requested_chunks,
            "successful": successful_chunks,
            "failed": failed_chunks,
        }

    minimum = min(values)
    maximum = max(values)

    add_weather_debug_log(
        debug_nonce,
        (
            f"500m 분할조회 완료 | min={minimum:.1f} | max={maximum:.1f} | "
            f"observations={len(values)} | "
            f"chunks={successful_chunks}/{requested_chunks}"
        ),
    )

    return {
        "available": True,
        "source": "500m 격자",
        "minimum": minimum,
        "maximum": maximum,
        "actual_end": actual_end.strftime("%H:%M"),
        "observations": len(values),
        "requested": requested_chunks,
        "successful": successful_chunks,
        "failed": failed_chunks,
    }


def latitude_longitude_to_grid(
    latitude: float,
    longitude: float,
) -> tuple[int, int]:
    """위·경도를 기상청 동네예보 Lambert 격자 좌표로 변환합니다."""
    earth_radius = 6371.00877
    grid_size = 5.0
    standard_parallel_1 = 30.0
    standard_parallel_2 = 60.0
    origin_longitude = 126.0
    origin_latitude = 38.0
    origin_x = 43.0
    origin_y = 136.0
    radians = math.pi / 180.0

    radius = earth_radius / grid_size
    parallel_1 = standard_parallel_1 * radians
    parallel_2 = standard_parallel_2 * radians
    origin_lon = origin_longitude * radians
    origin_lat = origin_latitude * radians

    sn = math.tan(math.pi * 0.25 + parallel_2 * 0.5) / math.tan(
        math.pi * 0.25 + parallel_1 * 0.5
    )
    sn = math.log(math.cos(parallel_1) / math.cos(parallel_2)) / math.log(sn)
    sf = math.tan(math.pi * 0.25 + parallel_1 * 0.5) ** sn
    sf = sf * math.cos(parallel_1) / sn
    ro = math.tan(math.pi * 0.25 + origin_lat * 0.5) ** sn
    ro = radius * sf / ro

    ra = math.tan(math.pi * 0.25 + latitude * radians * 0.5) ** sn
    ra = radius * sf / ra
    theta = longitude * radians - origin_lon
    if theta > math.pi:
        theta -= 2.0 * math.pi
    if theta < -math.pi:
        theta += 2.0 * math.pi
    theta *= sn

    grid_x = int(ra * math.sin(theta) + origin_x + 0.5)
    grid_y = int(ro - ra * math.cos(theta) + origin_y + 0.5)
    return grid_x, grid_y


def calculate_summer_apparent_temperature(
    air_temperature: float,
    relative_humidity: float,
) -> float:
    """기상청 여름철 산식으로 기온·습도 기반 체감온도를 계산합니다."""
    wet_bulb = (
        air_temperature
        * math.atan(0.151977 * math.sqrt(relative_humidity + 8.313659))
        + math.atan(air_temperature + relative_humidity)
        - math.atan(relative_humidity - 1.67633)
        + 0.00391838
        * relative_humidity ** 1.5
        * math.atan(0.023101 * relative_humidity)
        - 4.686035
    )
    apparent = (
        -0.2442
        + 0.55399 * wet_bulb
        + 0.45535 * air_temperature
        - 0.0022 * wet_bulb**2
        + 0.00278 * wet_bulb * air_temperature
        + 3.0
    )
    return round(apparent, 1)


def latest_village_forecast_base(now: datetime) -> datetime:
    """현재 시각에 이미 발표된 최신 단기예보 기준시각을 구합니다."""
    base_hours = (2, 5, 8, 11, 14, 17, 20, 23)
    # 단기예보는 기준시각 약 10분 뒤 제공되므로 15분의 여유를 둡니다.
    available_before = now - timedelta(minutes=15)
    candidates = [
        available_before.replace(hour=hour, minute=0, second=0, microsecond=0)
        for hour in base_hours
        if hour <= available_before.hour
    ]
    if candidates:
        return candidates[-1]
    previous_day = available_before - timedelta(days=1)
    return previous_day.replace(hour=23, minute=0, second=0, microsecond=0)


def fetch_kma_forecast_apparent_temperature_range(
    latitude: float,
    longitude: float,
    auth_key: str,
    range_start: datetime,
    range_end: datetime,
) -> tuple[str, str, str | None, str | None, int]:
    """단기예보로 근무시간의 예상 체감온도와 예상 폭염시간을 구합니다."""
    grid_x, grid_y = latitude_longitude_to_grid(latitude, longitude)
    base_time = latest_village_forecast_base(datetime.now(KST))
    params = urllib.parse.urlencode(
        {
            "pageNo": "1",
            "numOfRows": "2000",
            "dataType": "JSON",
            "base_date": base_time.strftime("%Y%m%d"),
            "base_time": base_time.strftime("%H%M"),
            "nx": str(grid_x),
            "ny": str(grid_y),
            "authKey": auth_key,
        }
    )
    request = urllib.request.Request(
        f"{KMA_VILLAGE_FORECAST_API_URL}?{params}",
        headers={"User-Agent": "checktemp-streamlit/1.0"},
    )

    with urllib.request.urlopen(request, timeout=20) as response:
        result = json.loads(response.read().decode("utf-8"))

    response_body = result.get("response") or {}
    header = response_body.get("header") or {}
    if clean_text(header.get("resultCode")) != "00":
        raise ValueError(
            clean_text(header.get("resultMsg")) or "단기예보 응답 오류"
        )

    raw_items = (
        response_body.get("body", {}).get("items", {}).get("item", [])
    )
    hourly_values: dict[datetime, dict[str, float]] = {}
    for item in raw_items:
        category = clean_text(item.get("category"))
        if category not in {"TMP", "REH"}:
            continue
        forecast_date = clean_text(item.get("fcstDate"))
        forecast_time = clean_text(item.get("fcstTime")).zfill(4)
        try:
            forecast_at = datetime.strptime(
                f"{forecast_date}{forecast_time}", "%Y%m%d%H%M"
            ).replace(tzinfo=KST)
        except ValueError:
            continue
        if not (range_start <= forecast_at <= range_end):
            continue
        value = parse_float(item.get("fcstValue"))
        if value is not None:
            hourly_values.setdefault(forecast_at, {})[category] = value

    forecasts: list[tuple[datetime, float]] = []
    for forecast_at, values in sorted(hourly_values.items()):
        air_temperature = values.get("TMP")
        relative_humidity = values.get("REH")
        if air_temperature is None or relative_humidity is None:
            continue
        if not (-50 <= air_temperature <= 60 and 0 <= relative_humidity <= 100):
            continue
        forecasts.append(
            (
                forecast_at,
                calculate_summer_apparent_temperature(
                    air_temperature,
                    relative_humidity,
                ),
            )
        )

    if not forecasts:
        raise ValueError(
            "선택한 날짜·근무시간은 현재 제공되는 단기예보 범위에 없습니다."
        )

    temperatures = [temperature for _, temperature in forecasts]
    heat_forecasts = [
        forecast_at
        for forecast_at, temperature in forecasts
        if temperature >= FORECAST_HEAT_THRESHOLD
    ]
    expected_start: str | None = None
    expected_end: str | None = None
    if heat_forecasts:
        heat_start = heat_forecasts[0]
        heat_end = min(heat_forecasts[-1] + timedelta(hours=1), range_end)
        expected_start = heat_start.strftime("%H:%M")
        expected_end = heat_end.strftime("%H:%M")

    return (
        format_number(min(temperatures)),
        format_number(max(temperatures)),
        expected_start,
        expected_end,
        len(forecasts),
    )


def fetch_kma_regional_apparent_temperature_at(
    latitude: float,
    longitude: float,
    auth_key: str,
    base_time: datetime,
    *,
    debug_nonce: int | None = None,
) -> tuple[float, float, float]:
    """지정 정시 초단기실황(T1H·REH)으로 체감온도를 계산합니다."""
    grid_x, grid_y = latitude_longitude_to_grid(latitude, longitude)
    base_time = base_time.astimezone(KST).replace(
        minute=0,
        second=0,
        microsecond=0,
    )

    params = urllib.parse.urlencode(
        {
            "pageNo": "1",
            "numOfRows": "1000",
            "dataType": "JSON",
            "base_date": base_time.strftime("%Y%m%d"),
            "base_time": base_time.strftime("%H%M"),
            "nx": str(grid_x),
            "ny": str(grid_y),
            "authKey": auth_key,
        }
    )

    request = urllib.request.Request(
        f"{KMA_ULTRA_NOWCAST_API_URL}?{params}",
        headers={"User-Agent": "checktemp-streamlit/1.0"},
    )

    started = time_module.perf_counter()

    try:
        with urllib.request.urlopen(
            request,
            timeout=KMA_REGIONAL_TIMEOUT_SECONDS,
        ) as response:
            result = json.loads(response.read().decode("utf-8"))
            elapsed = time_module.perf_counter() - started
            status = getattr(response, "status", None) or response.getcode()

        add_weather_debug_log(
            debug_nonce,
            (
                f"지역실황 HTTP 성공 | base={base_time.strftime('%Y%m%d %H:%M')} | "
                f"status={status} | elapsed={elapsed:.2f}s"
            ),
        )

    except Exception as error:
        elapsed = time_module.perf_counter() - started
        add_weather_debug_log(
            debug_nonce,
            (
                f"지역실황 HTTP 실패 | base={base_time.strftime('%Y%m%d %H:%M')} | "
                f"type={type(error).__name__} | elapsed={elapsed:.2f}s | "
                f"message={error}"
            ),
            level="error",
        )
        raise

    response_body = result.get("response") or {}
    header = response_body.get("header") or {}

    if clean_text(header.get("resultCode")) != "00":
        result_message = (
            clean_text(header.get("resultMsg"))
            or "초단기실황 응답 오류"
        )
        raise ValueError(result_message)

    raw_items = (
        response_body.get("body", {})
        .get("items", {})
        .get("item", [])
    )

    values = {
        clean_text(item.get("category")): parse_float(item.get("obsrValue"))
        for item in raw_items
    }

    air_temperature = values.get("T1H")
    relative_humidity = values.get("REH")

    if air_temperature is None or relative_humidity is None:
        raise ValueError("초단기실황에 기온 또는 습도 값이 없습니다.")

    if not (-50 <= air_temperature <= 60):
        raise ValueError("초단기실황 기온 값이 정상 범위를 벗어났습니다.")

    if not (0 <= relative_humidity <= 100):
        raise ValueError("초단기실황 습도 값이 정상 범위를 벗어났습니다.")

    apparent = calculate_summer_apparent_temperature(
        air_temperature,
        relative_humidity,
    )

    add_weather_debug_log(
        debug_nonce,
        (
            f"지역실황 값 | {base_time.strftime('%H:%M')} | "
            f"T1H={air_temperature:.1f} | REH={relative_humidity:.1f} | "
            f"apparent={apparent:.1f}"
        ),
    )

    return apparent, air_temperature, relative_humidity


def fetch_kma_regional_apparent_temperature_range(
    latitude: float,
    longitude: float,
    auth_key: str,
    range_start: datetime,
    range_end: datetime,
    *,
    debug_nonce: int | None = None,
) -> dict[str, Any]:
    """근무시간 안의 정시별 지역실황을 독립 조회해 체감온도 범위를 계산합니다."""
    now = datetime.now(KST).replace(second=0, microsecond=0)

    if range_start > now:
        raise ValueError("조회 시작시간이 현재보다 이후입니다.")

    actual_end = min(range_end, now)
    if actual_end < range_start:
        raise ValueError("조회할 시간대가 없습니다.")

    latest_available = now.replace(
        minute=0,
        second=0,
        microsecond=0,
    )
    if now.minute < 15:
        latest_available -= timedelta(hours=1)

    # 시작 이후 첫 정시부터 사용합니다.
    first_hour = range_start.replace(
        minute=0,
        second=0,
        microsecond=0,
    )
    if range_start.minute > 0:
        first_hour += timedelta(hours=1)

    last_hour = min(
        actual_end.replace(
            minute=0,
            second=0,
            microsecond=0,
        ),
        latest_available,
    )

    # 1시간보다 짧고 범위 안에 정시가 하나도 없으면
    # 시작 직전 정시를 대표값으로 1건 사용합니다.
    if first_hour > last_hour:
        fallback_hour = min(
            range_start.replace(
                minute=0,
                second=0,
                microsecond=0,
            ),
            latest_available,
        )
        sample_times = [fallback_hour]
    else:
        sample_times = []
        cursor = first_hour
        while cursor <= last_hour:
            sample_times.append(cursor)
            cursor += timedelta(hours=1)

    add_weather_debug_log(
        debug_nonce,
        (
            "지역실황 범위조회 시작 | "
            f"samples={len(sample_times)} | "
            f"from={sample_times[0].strftime('%Y-%m-%d %H:%M')} | "
            f"to={sample_times[-1].strftime('%Y-%m-%d %H:%M')}"
        ),
    )

    apparent_values: list[float] = []
    failed_hours: list[str] = []
    consecutive_failures = 0

    for sample_time in sample_times:
        label = sample_time.strftime("%H:%M")
        try:
            apparent, _, _ = fetch_kma_regional_apparent_temperature_at(
                latitude,
                longitude,
                auth_key,
                sample_time,
                debug_nonce=debug_nonce,
            )
            apparent_values.append(apparent)
            consecutive_failures = 0

        except Exception:
            failed_hours.append(label)
            consecutive_failures += 1

            # 지역실황 서버 자체가 연속으로 응답하지 않을 때도
            # 지나치게 오래 대기하지 않습니다.
            if (
                consecutive_failures >= KMA_REGIONAL_MAX_INITIAL_FAILURES
                and not apparent_values
            ):
                add_weather_debug_log(
                    debug_nonce,
                    (
                        "지역실황 회로차단 | "
                        f"첫 {KMA_REGIONAL_MAX_INITIAL_FAILURES}개 시간 "
                        "연속 실패로 "
                        "나머지 시간 조회 중단"
                    ),
                    level="warning",
                )
                break

    if not apparent_values:
        add_weather_debug_log(
            debug_nonce,
            (
                f"지역실황 범위조회 결과 없음 | "
                f"requested={len(sample_times)} | failed={len(failed_hours)}"
            ),
            level="warning",
        )
        return {
            "available": False,
            "source": "초단기실황",
            "minimum": None,
            "maximum": None,
            "actual_end": actual_end.strftime("%H:%M"),
            "observations": 0,
            "requested": len(sample_times),
            "successful": 0,
            "failed": failed_hours,
        }

    minimum = min(apparent_values)
    maximum = max(apparent_values)

    add_weather_debug_log(
        debug_nonce,
        (
            f"지역실황 범위조회 완료 | min={minimum:.1f} | "
            f"max={maximum:.1f} | "
            f"hours={len(apparent_values)}/{len(sample_times)}"
        ),
    )

    return {
        "available": True,
        "source": "초단기실황",
        "minimum": minimum,
        "maximum": maximum,
        "actual_end": actual_end.strftime("%H:%M"),
        "observations": len(apparent_values),
        "requested": len(sample_times),
        "successful": len(apparent_values),
        "failed": failed_hours,
    }


def fetch_open_meteo_archive_apparent_temperature_range(
    latitude: float,
    longitude: float,
    range_start: datetime,
    range_end: datetime,
    *,
    debug_nonce: int | None = None,
) -> dict[str, Any]:
    """기상청 조회가 모두 실패했을 때 쓰는 과거 체감온도 보조 조회입니다."""
    actual_end = min(range_end, datetime.now(KST).replace(second=0, microsecond=0))
    if actual_end < range_start:
        raise ValueError("조회할 시간대가 없습니다.")

    params = urllib.parse.urlencode(
        {
            "latitude": f"{latitude:.6f}",
            "longitude": f"{longitude:.6f}",
            "start_date": range_start.strftime("%Y-%m-%d"),
            "end_date": actual_end.strftime("%Y-%m-%d"),
            "hourly": "apparent_temperature",
            "timezone": "Asia/Seoul",
            "cell_selection": "nearest",
        }
    )
    request = urllib.request.Request(
        f"{OPEN_METEO_ARCHIVE_API_URL}?{params}",
        headers={"User-Agent": "checktemp-streamlit/1.0"},
    )

    add_weather_debug_log(
        debug_nonce,
        (
            "Open-Meteo 보조조회 시작 | "
            f"range={range_start.strftime('%Y-%m-%d %H:%M')}"
            f"~{actual_end.strftime('%Y-%m-%d %H:%M')} | "
            f"timeout={OPEN_METEO_ARCHIVE_TIMEOUT_SECONDS}s"
        ),
    )

    started = time_module.perf_counter()
    with urllib.request.urlopen(
        request,
        timeout=OPEN_METEO_ARCHIVE_TIMEOUT_SECONDS,
    ) as response:
        result = json.loads(response.read().decode("utf-8"))
        elapsed = time_module.perf_counter() - started
        status = getattr(response, "status", None) or response.getcode()

    hourly = result.get("hourly") or {}
    times = hourly.get("time") or []
    values = hourly.get("apparent_temperature") or []
    observations: list[float] = []

    for timestamp, value in zip(times, values):
        if value is None:
            continue
        try:
            observed_at = datetime.strptime(timestamp, "%Y-%m-%dT%H:%M").replace(
                tzinfo=KST
            )
        except ValueError:
            continue
        if not (range_start <= observed_at <= actual_end):
            continue
        temperature = parse_float(value)
        if temperature is not None:
            observations.append(temperature)

    if not observations:
        add_weather_debug_log(
            debug_nonce,
            (
                "Open-Meteo 보조조회 결과 없음 | "
                f"status={status} | elapsed={elapsed:.2f}s"
            ),
            level="warning",
        )
        return {
            "available": False,
            "source": "Open-Meteo 재분석",
            "minimum": None,
            "maximum": None,
            "actual_end": actual_end.strftime("%H:%M"),
            "observations": 0,
            "requested": len(times),
            "successful": 0,
            "failed": [],
        }

    minimum = min(observations)
    maximum = max(observations)

    add_weather_debug_log(
        debug_nonce,
        (
            f"Open-Meteo 보조조회 완료 | min={minimum:.1f} | "
            f"max={maximum:.1f} | observations={len(observations)} | "
            f"status={status} | elapsed={elapsed:.2f}s"
        ),
    )

    return {
        "available": True,
        "source": "Open-Meteo 재분석",
        "minimum": minimum,
        "maximum": maximum,
        "actual_end": actual_end.strftime("%H:%M"),
        "observations": len(observations),
        "requested": len(times),
        "successful": len(observations),
        "failed": [],
    }


def fetch_kma_regional_apparent_temperature(
    latitude: float,
    longitude: float,
    auth_key: str,
) -> tuple[str, str, str, str]:
    """동네예보 초단기실황의 기온·습도로 지역 체감온도를 계산합니다."""
    grid_x, grid_y = latitude_longitude_to_grid(latitude, longitude)
    current_time = datetime.now(KST).replace(second=0, microsecond=0)
    latest_candidate = current_time.replace(minute=0)
    if current_time.minute < 15:
        latest_candidate -= timedelta(hours=1)

    last_error = "조회 가능한 초단기실황이 없습니다."
    for hours_back in range(4):
        base_time = latest_candidate - timedelta(hours=hours_back)
        params = urllib.parse.urlencode(
            {
                "pageNo": "1",
                "numOfRows": "1000",
                "dataType": "JSON",
                "base_date": base_time.strftime("%Y%m%d"),
                "base_time": base_time.strftime("%H%M"),
                "nx": str(grid_x),
                "ny": str(grid_y),
                "authKey": auth_key,
            }
        )
        request = urllib.request.Request(
            f"{KMA_ULTRA_NOWCAST_API_URL}?{params}",
            headers={"User-Agent": "checktemp-streamlit/1.0"},
        )

        try:
            with urllib.request.urlopen(request, timeout=8) as response:
                result = json.loads(response.read().decode("utf-8"))
            response_body = result.get("response") or {}
            header = response_body.get("header") or {}
            if clean_text(header.get("resultCode")) != "00":
                last_error = clean_text(header.get("resultMsg")) or last_error
                continue

            raw_items = (
                response_body.get("body", {}).get("items", {}).get("item", [])
            )
            values = {
                clean_text(item.get("category")): parse_float(item.get("obsrValue"))
                for item in raw_items
            }
            air_temperature = values.get("T1H")
            relative_humidity = values.get("REH")
            if air_temperature is None or relative_humidity is None:
                last_error = "초단기실황에 기온 또는 습도 값이 없습니다."
                continue
            if not (-50 <= air_temperature <= 60 and 0 <= relative_humidity <= 100):
                last_error = "초단기실황 기온·습도 값이 정상 범위를 벗어났습니다."
                continue

            apparent = calculate_summer_apparent_temperature(
                air_temperature,
                relative_humidity,
            )
            return (
                format_number(apparent),
                base_time.strftime("%H:%M"),
                format_number(air_temperature),
                format_number(relative_humidity),
            )
        except Exception as error:  # noqa: BLE001
            last_error = str(error)

    raise ValueError(last_error)


def record_heat_start_with_weather(
    nonce: int,
    record_start_time: bool = True,
) -> None:
    """폭염 시작시간을 기록하고 설정 시 체감온도도 자동 입력합니다."""
    clear_weather_debug_log(nonce)
    add_weather_debug_log(
        nonce,
        f"조회 트리거 | record_start_time={record_start_time}",
    )
    if record_start_time:
        set_time_now("heat_start", nonce)
    notice_key = weather_notice_key(nonce)
    site_name = clean_text(st.session_state.get(f"site_{nonce}"))

    if not site_name:
        st.session_state[notice_key] = (
            "warning",
            "근무장소를 먼저 입력하면 체감온도를 자동 조회할 수 있습니다.",
        )
        return

    auth_key = get_secret(("weather", "kma_auth_key"))
    if not auth_key:
        st.session_state[notice_key] = (
            "info",
            "기상청 인증키가 아직 설정되지 않아 시간만 기록했습니다.",
        )
        return

    coordinates = get_site_coordinates(site_name)
    matched_name = site_name
    if coordinates is None:
        kakao_key = get_secret(("location", "kakao_rest_api_key"))
        if not kakao_key:
            st.session_state[notice_key] = (
                "warning",
                "현장 좌표가 등록되지 않았고 카카오 REST API 키도 없습니다. "
                "폭염 시작시간만 기록했습니다.",
            )
            return

        try:
            latitude, longitude, matched_name = (
                search_kakao_site_coordinates(site_name, kakao_key)
            )
            coordinates = (latitude, longitude)
        except Exception as error:  # noqa: BLE001
            st.session_state[notice_key] = (
                "warning",
                f"'{site_name}' 장소 검색에 실패해 시작시간만 기록했습니다. "
                f"근무장소를 더 정확하게 입력해 주세요. ({error})",
            )
            return

    heat_start_text = clean_text(
        st.session_state.get(f"manual_heat_start_{nonce}")
    )
    heat_end_text = clean_text(
        st.session_state.get(f"manual_heat_end_{nonce}")
    )
    work_start_text = clean_text(
        st.session_state.get(f"manual_work_start_{nonce}")
    )
    work_end_text = clean_text(
        st.session_state.get(f"manual_work_end_{nonce}")
    )

    selected_date = st.session_state.get(f"date_{nonce}")
    if not isinstance(selected_date, date):
        selected_date = datetime.now(KST).date()

    add_weather_debug_log(
        nonce,
        (
            f"입력값 | date={selected_date.isoformat()} | site={site_name} | "
            f"matched_site={matched_name} | "
            f"work={work_start_text or '-'}~{work_end_text or '-'} | "
            f"heat={heat_start_text or '-'}~{heat_end_text or '-'} | "
            f"coords={coordinates[0]:.6f},{coordinates[1]:.6f}"
        ),
    )

    if selected_date > datetime.now(KST).date():
        work_start_value = parse_time_value(work_start_text)
        work_end_value = parse_time_value(work_end_text)
        if work_start_value is None or work_end_value is None:
            st.session_state[notice_key] = (
                "warning",
                "미래 날짜의 예상 폭염시간을 조회하려면 근무 시작과 종료를 "
                "HH:MM 형식으로 모두 입력해 주세요.",
            )
            return

        forecast_start = datetime.combine(
            selected_date,
            work_start_value,
            tzinfo=KST,
        )
        forecast_end = datetime.combine(
            selected_date,
            work_end_value,
            tzinfo=KST,
        )
        if forecast_end < forecast_start:
            forecast_end += timedelta(days=1)

        try:
            (
                minimum,
                maximum,
                expected_start,
                expected_end,
                forecast_count,
            ) = fetch_kma_forecast_apparent_temperature_range(
                coordinates[0],
                coordinates[1],
                auth_key,
                forecast_start,
                forecast_end,
            )
        except Exception as error:  # noqa: BLE001
            if isinstance(error, urllib.error.HTTPError) and error.code == 403:
                error_detail = (
                    "현재 인증키에 '4.3 단기예보조회' 권한이 없습니다. "
                    "기상청 API허브에서 해당 API를 활용신청한 뒤 다시 조회해 "
                    "주세요. 기존에 승인된 2.2 초단기예보는 약 6시간 범위라 "
                    "미래 날짜 전체 조회에는 사용할 수 없습니다."
                )
            else:
                error_detail = (
                    "날짜가 단기예보 범위인지 확인해 주세요. "
                    f"({error})"
                )
            st.session_state[notice_key] = (
                "warning",
                "선택한 미래 날짜의 예상 체감온도 조회에 실패했습니다. "
                f"{error_detail}",
            )
            return

        temperature_range = (
            minimum if minimum == maximum else f"{minimum}~{maximum}"
        )
        st.session_state[f"temperature_{nonce}"] = temperature_range
        st.session_state[f"manual_heat_start_{nonce}"] = expected_start or ""
        st.session_state[f"manual_heat_end_{nonce}"] = expected_end or ""

        if expected_start and expected_end:
            heat_message = (
                f"체감온도 {format_number(FORECAST_HEAT_THRESHOLD)}℃ 이상 예상시간 "
                f"{expected_start}~{expected_end}를 예상 폭염시간으로 입력했습니다."
            )
        else:
            heat_message = (
                f"근무시간 중 체감온도 {format_number(FORECAST_HEAT_THRESHOLD)}℃ "
                "이상으로 예상되는 시간이 없어 폭염시간을 비워 두었습니다."
            )
        st.session_state[notice_key] = (
            "success",
            f"{matched_name} · {selected_date.strftime('%Y-%m-%d')} 단기예보 "
            f"체감온도 최저 {minimum}℃ · 최고 {maximum}℃를 "
            f"{forecast_count}개 시간자료로 확인했습니다. {heat_message} "
            "예보는 변경될 수 있으므로 당일 또는 작업 후 다시 조회해 주세요.",
        )
        save_weather_result_cache(nonce)
        return

    range_label = ""
    range_start_text = ""
    range_end_text = ""
    if heat_start_text and heat_end_text:
        range_label = "폭염시간"
        range_start_text = heat_start_text
        range_end_text = heat_end_text
    elif heat_start_text or heat_end_text:
        st.session_state[notice_key] = (
            "warning",
            "폭염시간을 사용하려면 폭염 시작과 종료를 모두 입력해 주세요.",
        )
        return
    elif work_start_text and work_end_text:
        range_label = "근무시간"
        range_start_text = work_start_text
        range_end_text = work_end_text
    elif work_start_text or work_end_text:
        st.session_state[notice_key] = (
            "warning",
            "폭염시간이 비어 있으면 근무시간을 사용합니다. 근무 시작과 "
            "종료를 모두 입력해 주세요.",
        )
        return

    if range_label:
        range_start_value = parse_time_value(range_start_text)
        range_end_value = parse_time_value(range_end_text)
        if range_start_value is None or range_end_value is None:
            st.session_state[notice_key] = (
                "warning",
                f"{range_label} 시작과 종료를 HH:MM 형식으로 입력해 주세요.",
            )
            return

        range_start = datetime.combine(
            selected_date,
            range_start_value,
            tzinfo=KST,
        )
        range_end = datetime.combine(
            selected_date,
            range_end_value,
            tzinfo=KST,
        )
        if range_end < range_start:
            range_end += timedelta(days=1)
            add_weather_debug_log(
                nonce,
                "자정 경과 인식 | 종료시간을 익일로 처리",
            )

        add_weather_debug_log(
            nonce,
            (
                f"{range_label} 해석 완료 | "
                f"start={range_start.strftime('%Y-%m-%d %H:%M')} | "
                f"end={range_end.strftime('%Y-%m-%d %H:%M')}"
            ),
        )

        grid_result: dict[str, Any] | None = None
        regional_result: dict[str, Any] | None = None
        range_errors: list[str] = []
        use_regional_fallback = selected_date >= datetime.now(KST).date()

        try:
            grid_result = fetch_kma_apparent_temperature_range(
                coordinates[0],
                coordinates[1],
                auth_key,
                range_start,
                range_end,
                debug_nonce=nonce,
            )
        except Exception as error:  # noqa: BLE001
            range_errors.append(f"500m 격자: {error}")
            add_weather_debug_log(
                nonce,
                (
                    f"500m 소스 예외 종료 | type={type(error).__name__} | "
                    f"message={error}"
                ),
                level="error",
            )

        if use_regional_fallback:
            try:
                regional_result = fetch_kma_regional_apparent_temperature_range(
                    coordinates[0],
                    coordinates[1],
                    auth_key,
                    range_start,
                    range_end,
                    debug_nonce=nonce,
                )
            except Exception as error:  # noqa: BLE001
                range_errors.append(f"초단기실황: {error}")
                add_weather_debug_log(
                    nonce,
                    (
                        f"지역실황 소스 예외 종료 | type={type(error).__name__} | "
                        f"message={error}"
                    ),
                    level="error",
                )
        else:
            add_weather_debug_log(
                nonce,
                (
                    "지역실황 생략 | 지난 날짜는 초단기실황 대신 "
                    "과거 재분석 보조조회로 진행"
                ),
                level="warning",
            )

        available_results = [
            result
            for result in (grid_result, regional_result)
            if result and result.get("available")
        ]

        archive_result: dict[str, Any] | None = None
        if not available_results:
            try:
                archive_result = (
                    fetch_open_meteo_archive_apparent_temperature_range(
                        coordinates[0],
                        coordinates[1],
                        range_start,
                        range_end,
                        debug_nonce=nonce,
                    )
                )
            except Exception as error:  # noqa: BLE001
                range_errors.append(f"Open-Meteo 재분석: {error}")
                add_weather_debug_log(
                    nonce,
                    (
                        "Open-Meteo 보조조회 예외 종료 | "
                        f"type={type(error).__name__} | message={error}"
                    ),
                    level="error",
                )

            if archive_result and archive_result.get("available"):
                available_results.append(archive_result)

        if not available_results:
            error_detail = (
                " / ".join(range_errors)
                if range_errors
                else "두 기상자료 모두 유효한 값을 받지 못했습니다."
            )
            st.session_state[notice_key] = (
                "warning",
                f"입력한 {range_label}의 체감온도 조회에 실패했습니다. "
                f"{error_detail} 기상조회 진단 로그를 확인해 주세요.",
            )
            return

        # 안전 우선: 최고 체감온도가 더 높은 소스의 전체 범위를 적용합니다.
        # 최고값이 같으면 최저값이 더 높은 결과를 우선합니다.
        selected_result = max(
            available_results,
            key=lambda result: (
                float(result["maximum"]),
                float(result["minimum"]),
            ),
        )

        selected_minimum = float(selected_result["minimum"])
        selected_maximum = float(selected_result["maximum"])

        minimum_text = format_number(selected_minimum)
        maximum_text = format_number(selected_maximum)
        temperature_range = (
            minimum_text
            if minimum_text == maximum_text
            else f"{minimum_text}~{maximum_text}"
        )

        st.session_state[f"temperature_{nonce}"] = temperature_range

        requested_end_text = range_end.strftime("%H:%M")
        actual_end_text = clean_text(selected_result.get("actual_end"))
        end_note = (
            ""
            if actual_end_text == requested_end_text
            else f" (현재 조회 가능한 {actual_end_text}까지)"
        )

        source_details: list[str] = []

        if grid_result:
            if grid_result.get("available"):
                grid_min = format_number(float(grid_result["minimum"]))
                grid_max = format_number(float(grid_result["maximum"]))
                grid_range = (
                    grid_min
                    if grid_min == grid_max
                    else f"{grid_min}~{grid_max}"
                )
                source_details.append(
                    "500m 격자 "
                    f"{grid_range}℃ "
                    f"({grid_result['successful']}/{grid_result['requested']}구간, "
                    f"{grid_result['observations']}개 관측)"
                )
            else:
                source_details.append("500m 격자 조회 실패")

        if regional_result:
            if regional_result.get("available"):
                regional_min = format_number(
                    float(regional_result["minimum"])
                )
                regional_max = format_number(
                    float(regional_result["maximum"])
                )
                regional_range = (
                    regional_min
                    if regional_min == regional_max
                    else f"{regional_min}~{regional_max}"
                )
                source_details.append(
                    "초단기실황 "
                    f"{regional_range}℃ "
                    f"({regional_result['successful']}/"
                    f"{regional_result['requested']}시간)"
                )
            else:
                source_details.append("초단기실황 조회 실패")

        if archive_result:
            if archive_result.get("available"):
                archive_min = format_number(float(archive_result["minimum"]))
                archive_max = format_number(float(archive_result["maximum"]))
                archive_range = (
                    archive_min
                    if archive_min == archive_max
                    else f"{archive_min}~{archive_max}"
                )
                source_details.append(
                    "Open-Meteo 재분석 "
                    f"{archive_range}℃ "
                    f"({archive_result['observations']}개 시간자료, "
                    "기상청 장애 시 보조값)"
                )
            else:
                source_details.append("Open-Meteo 재분석 조회 실패")

        if range_errors:
            source_details.append(
                f"일부 소스 오류: {' / '.join(range_errors)}"
            )

        selected_source = clean_text(selected_result.get("source"))

        add_weather_debug_log(
            nonce,
            (
                f"안전 우선 최종선택 | source={selected_source} | "
                f"range={temperature_range} | "
                f"max={maximum_text}"
            ),
        )

        st.session_state[notice_key] = (
            "success",
            f"{matched_name} · {range_label} "
            f"{range_start.strftime('%H:%M')}~"
            f"{requested_end_text}{end_note} · "
            f"{' / '.join(source_details)} · "
            f"안전을 위해 최고 체감온도가 더 높은 "
            f"{selected_source} 결과 {temperature_range}℃를 자동 적용했습니다.",
        )
        save_weather_result_cache(nonce)
        return

    grid_temperature: str | None = None
    grid_observed_at = ""
    regional_temperature: str | None = None
    regional_observed_at = ""
    regional_air_temperature = ""
    regional_humidity = ""
    lookup_errors: list[str] = []

    try:
        grid_temperature, grid_observed_at = fetch_kma_apparent_temperature(
            coordinates[0],
            coordinates[1],
            auth_key,
        )
    except Exception as error:  # noqa: BLE001
        lookup_errors.append(f"500m 격자: {error}")

    try:
        (
            regional_temperature,
            regional_observed_at,
            regional_air_temperature,
            regional_humidity,
        ) = fetch_kma_regional_apparent_temperature(
            coordinates[0],
            coordinates[1],
            auth_key,
        )
    except Exception as error:  # noqa: BLE001
        lookup_errors.append(f"지역 실황: {error}")

    available_temperatures = [
        value
        for value in (grid_temperature, regional_temperature)
        if value is not None
    ]
    if not available_temperatures:
        st.session_state[notice_key] = (
            "warning",
            "기상청 체감온도 조회에 실패해 시간만 기록했습니다. "
            f"직접 입력해 주세요. ({' / '.join(lookup_errors)})",
        )
        return

    applied_temperature = max(
        available_temperatures,
        key=lambda value: float(value),
    )
    st.session_state[f"temperature_{nonce}"] = applied_temperature

    detail_parts: list[str] = []
    if grid_temperature is not None:
        detail_parts.append(
            f"500m 격자 {grid_observed_at} {grid_temperature}℃"
        )
    if regional_temperature is not None:
        detail_parts.append(
            f"지역 실황 {regional_observed_at} {regional_temperature}℃"
            f"(기온 {regional_air_temperature}℃·습도 {regional_humidity}%)"
        )
    if lookup_errors:
        detail_parts.append(f"일부 조회 실패: {' / '.join(lookup_errors)}")

    st.session_state[notice_key] = (
        "success",
        f"{matched_name} · {' / '.join(detail_parts)} · "
        f"안전을 위해 높은 값 {applied_temperature}℃를 자동 적용했습니다.",
    )
    save_weather_result_cache(nonce)


def auto_lookup_future_weather(nonce: int) -> None:
    selected_date = st.session_state.get(f"date_{nonce}")
    site_name = clean_text(st.session_state.get(f"site_{nonce}"))
    work_start = clean_text(
        st.session_state.get(f"manual_work_start_{nonce}")
    )
    work_end = clean_text(
        st.session_state.get(f"manual_work_end_{nonce}")
    )

    if (
        not isinstance(selected_date, date)
        or selected_date <= datetime.now(KST).date()
        or not site_name
        or parse_time_value(work_start) is None
        or parse_time_value(work_end) is None
    ):
        return

    request_signature = (
        selected_date.isoformat(),
        site_name,
        work_start,
        work_end,
    )
    signature_key = f"auto_forecast_signature_{nonce}"
    if st.session_state.get(signature_key) == request_signature:
        return

    st.session_state[signature_key] = request_signature
    record_heat_start_with_weather(nonce, False)


def normalize_manual_time_field(
    field: str,
    nonce: int,
    *,
    trigger_weather: bool = False,
) -> None:
    """직접 입력한 시간을 HH:MM으로 정규화하고 필요 시 예보 조회를 실행합니다.

    예: 2256 -> 22:56, 930 -> 09:30, 18 -> 18:00
    """
    key = f"manual_{field}_{nonce}"
    raw_value = clean_text(st.session_state.get(key))

    if not raw_value:
        if trigger_weather:
            auto_lookup_future_weather(nonce)
        return

    parsed = parse_time_value(raw_value)
    if parsed is None:
        # 잘못된 값은 사용자가 수정할 수 있도록 그대로 둡니다.
        return

    # callback 단계에서 widget state를 먼저 정규화하므로
    # Enter 후 화면에도 즉시 HH:MM 형식으로 표시됩니다.
    st.session_state[key] = format_time_value(parsed)

    if trigger_weather:
        auto_lookup_future_weather(nonce)


def clear_time(field: str, nonce: int) -> None:
    st.session_state[time_state_key(field, nonce)] = None


def time_state_text(field: str, nonce: int) -> str:
    return format_time_value(
        st.session_state.get(time_state_key(field, nonce))
    )


def team_state_key(nonce: int) -> str:
    return f"selected_team_{nonce}"


def initialize_team_state(
    editing_record: dict[str, Any],
    nonce: int,
) -> None:
    key = team_state_key(nonce)
    if key not in st.session_state:
        current_team = clean_text(editing_record.get("팀"))
        current_team = {
            "중계팀": "중계",
            "영상팀": "영상",
        }.get(current_team, current_team)
        st.session_state[key] = (
            current_team
            if current_team in TEAM_OPTIONS
            else ""
        )


def set_team(team: str, nonce: int) -> None:
    if team in TEAM_OPTIONS:
        st.session_state[team_state_key(nonce)] = team


def sync_worker_count(nonce: int) -> None:
    employee_count = parse_int(
        st.session_state.get(f"employee_count_{nonce}")
    )
    contractor_count = parse_int(
        st.session_state.get(f"contractor_count_{nonce}")
    )
    st.session_state[f"worker_count_{nonce}"] = (
        employee_count + contractor_count
    )


def measures_state_key(nonce: int) -> str:
    return f"selected_measures_{nonce}"


def initialize_measures_state(
    editing_record: dict[str, Any],
    nonce: int,
) -> None:
    key = measures_state_key(nonce)
    if key not in st.session_state:
        st.session_state[key] = selected_measures(
            editing_record.get("조치사항")
        )


def toggle_measure(measure: str, nonce: int) -> None:
    key = measures_state_key(nonce)
    current = list(st.session_state.get(key, []))

    if measure in current:
        current.remove(measure)
    else:
        current.append(measure)

    # 화면/저장 순서는 MEASURE_OPTIONS 순서로 고정
    st.session_state[key] = [
        option for option in MEASURE_OPTIONS if option in current
    ]


def rest_minutes_key(nonce: int) -> str:
    return f"rest_minutes_{nonce}"


def initialize_rest_minutes(
    editing_record: dict[str, Any],
    nonce: int,
) -> None:
    key = rest_minutes_key(nonce)
    if key not in st.session_state:
        st.session_state[key] = max(
            0,
            parse_int(editing_record.get("휴게시간"), 0),
        )


def add_rest_minutes(minutes: int, nonce: int) -> None:
    key = rest_minutes_key(nonce)
    current = parse_int(st.session_state.get(key), 0)
    st.session_state[key] = max(0, current + minutes)


def clear_rest_minutes(nonce: int) -> None:
    st.session_state[rest_minutes_key(nonce)] = 0


def calculate_minutes(start: str, end: str) -> int:
    if not start or not end:
        return 0

    start_hour, start_minute = map(int, start.split(":"))
    end_hour, end_minute = map(int, end.split(":"))

    minutes = (
        end_hour * 60
        + end_minute
        - start_hour * 60
        - start_minute
    )
    return minutes + 1440 if minutes < 0 else minutes


def render_time_summary(nonce: int) -> None:
    work_start = time_state_text("work_start", nonce)
    work_end = time_state_text("work_end", nonce)
    heat_start = time_state_text("heat_start", nonce)
    heat_end = time_state_text("heat_end", nonce)
    rest_minutes = parse_int(
        st.session_state.get(rest_minutes_key(nonce)),
        0,
    )

    work_text = f"{work_start or '-'} ~ {work_end or '-'}"
    heat_text = f"{heat_start or '-'} ~ {heat_end or '-'}"
    rest_text = f"{rest_minutes}분" if rest_minutes > 0 else "-"

    st.markdown(
        f"""
        <div class="quick-time-summary">
            <div class="quick-time-card">
                <span>근무 시간</span>
                <strong>{html.escape(work_text)}</strong>
            </div>
            <div class="quick-time-card">
                <span>폭염 노출</span>
                <strong>{html.escape(heat_text)}</strong>
            </div>
            <div class="quick-time-card">
                <span>누적 휴게시간</span>
                <strong>{html.escape(rest_text)}</strong>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


@st.cache_resource(show_spinner=False)
def get_worksheet() -> gspread.Worksheet:
    service_account = dict(st.secrets["google_service_account"])

    if "private_key" in service_account:
        service_account["private_key"] = str(
            service_account["private_key"]
        ).replace("\\n", "\n")

    spreadsheet_url = get_secret(("app", "spreadsheet_url"))
    worksheet_name = get_secret(
        ("app", "worksheet"),
        WORKSHEET_DEFAULT,
    )

    if not spreadsheet_url:
        raise ValueError(
            "Streamlit Secrets의 app.spreadsheet_url 값이 비어 있습니다."
        )

    client = gspread.service_account_from_dict(service_account)
    spreadsheet = client.open_by_url(spreadsheet_url)
    return spreadsheet.worksheet(worksheet_name)


def column_letter(index: int) -> str:
    result = ""
    value = index
    while value:
        value, remainder = divmod(value - 1, 26)
        result = chr(65 + remainder) + result
    return result


def infer_sport(site_name: Any) -> str:
    site = clean_text(site_name)
    if any(keyword in site for keyword in ("야구", "구장", "베이스볼")):
        return "프로야구"
    if any(keyword in site.upper() for keyword in ("KLPGA", "여자골프")):
        return "KLPGA"
    if (
        "KPGA" in site.upper()
        or any(keyword in site for keyword in ("골프", "CC", "cc", "컨트리클럽"))
    ):
        return "KPGA"
    return "기타/ENG(실내)"


def normalize_sport(value: Any, site_name: Any = "") -> str:
    sport = clean_text(value)
    legacy_map = {
        "야구": "프로야구",
        "남자골프": "KPGA",
        "여자골프": "KLPGA",
        "기타스포츠(실외)": "기타/ENG(실외)",
        "기타스포츠(실내)": "기타/ENG(실내)",
        "기타스포츠": "기타/ENG(실내)",
        "기타": "기타/ENG(실내)",
        "기타 (실외)": "기타/ENG(실외)",
        "기타 (실내)": "기타/ENG(실내)",
    }
    if sport in legacy_map:
        return legacy_map[sport]
    if sport == "골프":
        return infer_sport(site_name) if site_name else "KPGA"
    if sport in SPORT_OPTIONS:
        return sport
    return infer_sport(site_name)


def common_measures_for_sport(sport: Any) -> str:
    normalized = normalize_sport(sport)
    return SPORT_COMMON_MEASURES.get(
        normalized,
        SPORT_COMMON_MEASURES["기타/ENG(실내)"],
    )


def strip_legacy_common_measures(sport: Any, value: Any) -> str:
    """기존 조치사항에서 v3.16 공통조치 문구만 제거합니다."""
    text = clean_text(value)
    if not text:
        return ""

    normalized_sport = normalize_sport(sport)
    legacy_key = {
        "프로야구": "야구",
        "KPGA": "남자골프",
        "KLPGA": "여자골프",
        "기타/ENG(실외)": "기타스포츠(실외)",
        "기타/ENG(실내)": "기타스포츠(실내)",
    }.get(normalized_sport, normalized_sport)
    legacy = LEGACY_COMMON_MEASURES.get(legacy_key, "")
    legacy_parts = {
        part.strip()
        for part in legacy.split("|")
        if part.strip()
    }
    parts = [
        part.strip()
        for part in text.split("|")
        if part.strip()
    ]
    return " | ".join(
        part for part in parts
        if part not in legacy_parts
    )


def migrate_sheet_to_current(
    worksheet: gspread.Worksheet,
    values: list[list[str]],
) -> None:
    backup_name = (
        f"{worksheet.title}_backup_"
        f"{datetime.now(KST).strftime('%Y%m%d_%H%M%S')}"
    )
    worksheet.duplicate(new_sheet_name=backup_name)

    old_headers = [clean_text(value) for value in values[0]]
    migrated_rows: list[list[str]] = [COLUMNS]

    for raw_row in values[1:]:
        padded = raw_row + [""] * max(
            0,
            len(old_headers) - len(raw_row),
        )
        old_record = {
            header: clean_text(padded[index])
            for index, header in enumerate(old_headers)
            if header
        }
        if not any(old_record.values()):
            continue

        sport = normalize_sport(
            old_record.get("종목"),
            old_record.get("현장명"),
        )
        new_record = {
            column: old_record.get(column, "")
            for column in COLUMNS
        }
        new_record["종목"] = sport
        new_record["팀"] = {
            "중계팀": "중계",
            "영상팀": "영상",
        }.get(clean_text(old_record.get("팀")), clean_text(old_record.get("팀")))
        new_record["공통 조치사항"] = (
            old_record.get("공통 조치사항")
            or common_measures_for_sport(sport)
        )
        new_record["조치사항"] = strip_legacy_common_measures(
            sport,
            old_record.get("조치사항"),
        )
        migrated_rows.append(record_values(new_record))

    worksheet.clear()
    worksheet.update(
        range_name="A1",
        values=migrated_rows,
        value_input_option="USER_ENTERED",
    )


def ensure_headers(worksheet: gspread.Worksheet) -> list[str]:
    headers = [
        clean_text(value)
        for value in worksheet.row_values(1)
    ]

    if not headers:
        worksheet.update(
            range_name=f"A1:{column_letter(len(COLUMNS))}1",
            values=[COLUMNS],
        )
        return COLUMNS

    if headers[: len(COLUMNS)] == COLUMNS:
        return headers

    if (
        headers[: len(OLD_CURRENT_COLUMNS)] == OLD_CURRENT_COLUMNS
        or
        headers[: len(PRE_COMMON_COLUMNS)] == PRE_COMMON_COLUMNS
        or headers[: len(LEGACY_COLUMNS)] == LEGACY_COLUMNS
    ):
        values = worksheet.get_all_values()
        migrate_sheet_to_current(worksheet, values)
        return COLUMNS

    missing = [
        column
        for column in COLUMNS
        if column not in headers
    ]
    missing_text = (
        ", ".join(missing)
        if missing
        else "열 순서 불일치"
    )
    raise ValueError(
        "records 시트의 첫 행 제목이 앱 형식과 다릅니다. "
        f"확인할 항목: {missing_text}"
    )


def _sheet_grid_borders(
    style: str = "SOLID",
) -> dict[str, Any]:
    """Google Sheets 표 전체 셀에 적용할 네 방향 테두리입니다."""
    color = {"red": 0.76, "green": 0.79, "blue": 0.84}
    return {
        edge: {
            "style": style,
            "color": color,
        }
        for edge in ("top", "bottom", "left", "right")
    }


def _data_row_format(row_number: int) -> dict[str, Any]:
    """데이터 행 구분용 교차 음영·전체 테두리·줄바꿈 서식입니다."""
    background = (
        {"red": 1.0, "green": 1.0, "blue": 1.0}
        if row_number % 2 == 0
        else {"red": 0.965, "green": 0.973, "blue": 0.984}
    )
    return {
        "backgroundColor": background,
        "textFormat": {
            "bold": False,
            "foregroundColor": {"red": 0.08, "green": 0.10, "blue": 0.14},
        },
        "verticalAlignment": "MIDDLE",
        "wrapStrategy": "WRAP",
        "borders": _sheet_grid_borders(),
    }


def _special_notes_format(row_number: int, has_notes: bool) -> dict[str, Any]:
    if has_notes:
        return {
            "backgroundColor": {"red": 1.0, "green": 0.93, "blue": 0.80},
            "textFormat": {
                "bold": True,
                "foregroundColor": {"red": 0.55, "green": 0.20, "blue": 0.05},
            },
            "verticalAlignment": "MIDDLE",
            "wrapStrategy": "WRAP",
            "borders": _sheet_grid_borders(),
        }

    base = _data_row_format(row_number)
    return {
        "backgroundColor": base["backgroundColor"],
        "textFormat": base["textFormat"],
        "verticalAlignment": "MIDDLE",
        "wrapStrategy": "WRAP",
        "borders": _sheet_grid_borders(),
    }


def format_special_notes_cell(
    worksheet: gspread.Worksheet,
    row_number: int,
    notes: Any,
) -> None:
    notes_column = column_letter(COLUMNS.index("특이사항") + 1)
    has_notes = bool(clean_text(notes))

    # 행 전체가 아니라 특이사항 셀 하나에만 강조를 적용합니다.
    worksheet.format(
        f"{notes_column}{row_number}",
        _special_notes_format(row_number, has_notes),
    )


def normalize_existing_special_notes_format(
    worksheet: gspread.Worksheet,
    values: list[list[str]],
) -> None:
    """기존 데이터 전체 서식 함수.

    앱 시작 시 호출하지 않습니다. 전체 행 서식은 느릴 수 있으므로
    신규/수정 행은 format_special_notes_cell()로 개별 적용합니다.
    """
    if not values:
        return

    session_key = (
        f"sheet_visual_v320_{worksheet.id}_{len(values)}_{len(COLUMNS)}"
    )
    if st.session_state.get(session_key):
        return

    headers = [clean_text(value) for value in values[0]]
    if "특이사항" not in headers:
        return

    notes_index = headers.index("특이사항")
    notes_column = column_letter(notes_index + 1)
    end_column = column_letter(len(COLUMNS))

    formats: list[dict[str, Any]] = [
        {
            "range": f"A1:{end_column}1",
            "format": {
                "backgroundColor": {"red": 0.09, "green": 0.17, "blue": 0.30},
                "textFormat": {
                    "bold": True,
                    "foregroundColor": {"red": 1.0, "green": 1.0, "blue": 1.0},
                },
                "horizontalAlignment": "CENTER",
                "verticalAlignment": "MIDDLE",
                "wrapStrategy": "WRAP",
                "borders": {
                    edge: {
                        "style": "SOLID_MEDIUM",
                        "color": {"red": 0.09, "green": 0.17, "blue": 0.30},
                    }
                    for edge in ("top", "bottom", "left", "right")
                },
            },
        }
    ]

    for row_number, raw_row in enumerate(values[1:], start=2):
        notes = raw_row[notes_index] if notes_index < len(raw_row) else ""
        if clean_text(notes):
            formats.append(
                {
                    "range": f"{notes_column}{row_number}",
                    "format": _special_notes_format(row_number, True),
                }
            )

    worksheet.batch_format(formats)

    try:
        worksheet.freeze(rows=1)
    except Exception:  # noqa: BLE001
        pass

    st.session_state[session_key] = True

@st.cache_data(ttl=20, show_spinner=False)
def load_records() -> pd.DataFrame:
    """records 시트를 1회만 읽어 앱 로딩을 가볍게 유지합니다.

    기존 전체 행 서식 재적용은 앱 시작 경로에서 제거했습니다.
    새 기록/수정 행은 저장 시 해당 행만 서식 적용합니다.
    """
    worksheet = get_worksheet()
    values = worksheet.get_all_values()

    if not values:
        worksheet.update(
            range_name=f"A1:{column_letter(len(COLUMNS))}1",
            values=[COLUMNS],
        )
        return empty_dataframe()

    headers = [clean_text(value) for value in values[0]]

    if headers[: len(COLUMNS)] == COLUMNS:
        pass
    elif (
        headers[: len(OLD_CURRENT_COLUMNS)] == OLD_CURRENT_COLUMNS
        or headers[: len(PRE_COMMON_COLUMNS)] == PRE_COMMON_COLUMNS
        or headers[: len(LEGACY_COLUMNS)] == LEGACY_COLUMNS
    ):
        migrate_sheet_to_current(worksheet, values)
        # 마이그레이션은 예외적인 1회 작업이므로 이후에만 다시 읽습니다.
        values = worksheet.get_all_values()
    else:
        missing = [
            column
            for column in COLUMNS
            if column not in headers
        ]
        missing_text = (
            ", ".join(missing)
            if missing
            else "열 순서 불일치"
        )
        raise ValueError(
            "records 시트의 첫 행 제목이 앱 형식과 다릅니다. "
            f"확인할 항목: {missing_text}"
        )

    if len(values) <= 1:
        return empty_dataframe()

    rows: list[dict[str, str]] = []

    for raw_row in values[1:]:
        padded = raw_row + [""] * (
            len(COLUMNS) - len(raw_row)
        )
        record = {
            column: clean_text(padded[index])
            for index, column in enumerate(COLUMNS)
        }

        if any(record.values()):
            rows.append(record)

    if not rows:
        return empty_dataframe()

    return pd.DataFrame(
        rows,
        columns=COLUMNS,
    ).fillna("")


def record_values(record: dict[str, Any]) -> list[str]:
    return [
        clean_text(record.get(column, ""))
        for column in COLUMNS
    ]


def append_record(record: dict[str, Any]) -> None:
    worksheet = get_worksheet()
    ensure_headers(worksheet)
    append_result = worksheet.append_row(
        record_values(record),
        value_input_option="USER_ENTERED",
        insert_data_option="INSERT_ROWS",
    )
    updated_range = clean_text(
        (append_result.get("updates") or {}).get("updatedRange")
        if isinstance(append_result, dict)
        else ""
    )
    row_match = re.search(r"(\d+)(?::[A-Z]+\d+)?$", updated_range)
    row_number = int(row_match.group(1)) if row_match else 0
    if row_number <= 0:
        appended_cell = worksheet.find(
            clean_text(record.get("id")),
            in_column=1,
        )
        row_number = appended_cell.row if appended_cell else 0
    try:
        if row_number > 0:
            format_special_notes_cell(
                worksheet,
                row_number,
                record.get("특이사항"),
            )
    except Exception:  # noqa: BLE001
        # 기록 저장은 성공했으므로 서식 오류 때문에 중복 저장되지 않게 합니다.
        pass

    load_records.clear()


def update_record(
    record_id: str,
    record: dict[str, Any],
) -> None:
    worksheet = get_worksheet()
    ensure_headers(worksheet)
    cell = worksheet.find(record_id, in_column=1)

    if cell is None:
        raise ValueError(
            "수정할 기록을 찾지 못했습니다. "
            "목록을 새로고침해 주세요."
        )

    worksheet.update(
        range_name=(
            f"A{cell.row}:"
            f"{column_letter(len(COLUMNS))}{cell.row}"
        ),
        values=[record_values(record)],
        value_input_option="USER_ENTERED",
    )
    try:
        format_special_notes_cell(
            worksheet,
            cell.row,
            record.get("특이사항"),
        )
    except Exception:  # noqa: BLE001
        pass

    load_records.clear()


def delete_record(record_id: str) -> None:
    worksheet = get_worksheet()
    ensure_headers(worksheet)
    cell = worksheet.find(record_id, in_column=1)

    if cell is None:
        raise ValueError(
            "삭제할 기록을 찾지 못했습니다. "
            "목록을 새로고침해 주세요."
        )

    worksheet.delete_rows(cell.row)
    load_records.clear()


def heat_level(value: Any) -> str:
    temperature = max_temperature(value)

    if temperature is None:
        return "온도 미입력"
    if temperature >= 38:
        return "매우 위험"
    if temperature >= 35:
        return "위험"
    if temperature >= 33:
        return "주의"
    if temperature >= 31:
        return "관심"
    return "일반"


def heat_level_class(value: Any) -> str:
    temperature = max_temperature(value)

    if temperature is None:
        return ""
    if temperature >= 35:
        return "status-danger"
    if temperature >= 31:
        return "status-caution"
    return "status-normal"


def option_index(
    options: list[str],
    value: Any,
    fallback: int = 0,
) -> int:
    text = (
        normalize_sport(value)
        if options == SPORT_OPTIONS
        else clean_text(value)
    )

    try:
        return options.index(text)
    except ValueError:
        return fallback


def selected_measures(value: Any) -> list[str]:
    text = clean_text(value)

    if not text:
        return []

    parts = [
        part.strip()
        for part in re.split(r"\s*\|\s*|,\s*", text)
        if part.strip()
    ]

    return [
        part
        for part in parts
        if part in MEASURE_OPTIONS
    ]


def custom_measure_text(value: Any) -> str:
    """Return saved action text that is not represented by a choice button."""
    text = clean_text(value)

    if not text:
        return ""

    parts = [
        part.strip()
        for part in re.split(r"\s*\|\s*|,\s*", text)
        if part.strip()
    ]
    return " | ".join(
        part for part in parts
        if part not in MEASURE_OPTIONS
        and not part.startswith(TIME_ADJUST_REASON_PREFIX)
    )


def time_adjust_reason_text(value: Any) -> str:
    parts = [
        part.strip()
        for part in re.split(r"\s*\|\s*", clean_text(value))
        if part.strip()
    ]
    for part in parts:
        if part.startswith(TIME_ADJUST_REASON_PREFIX):
            return part[len(TIME_ADJUST_REASON_PREFIX):].strip()
    return ""


def make_csv_bytes(dataframe: pd.DataFrame) -> bytes:
    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")
    writer.writerow(
        [
            {
                "작업날짜": "근무일자",
                "종목": "업무내용",
                "현장명": "근무장소",
                "팀": "부서",
            }.get(column, column)
            for column in COLUMNS
        ]
    )

    for _, row in dataframe[COLUMNS].iterrows():
        writer.writerow(
            [clean_text(row[column]) for column in COLUMNS]
        )

    return output.getvalue().encode("utf-8-sig")


def unique_texts(values: list[Any]) -> list[str]:
    results: list[str] = []
    for value in values:
        text = clean_text(value)
        if text and text not in results:
            results.append(text)
    return results


def report_team_summary(records: pd.DataFrame, team: str) -> str:
    aliases = {team, team.removesuffix("팀")}
    team_rows = records[records["팀"].isin(aliases)]
    if team_rows.empty:
        return "- 기록 없음"

    measures = unique_texts(team_rows["조치사항"].tolist())
    cleaned_measures: list[str] = []
    legacy_camera_pattern = re.compile(
        r"(?:[·ㆍ-]?\s*필드\s*카메라\s*[:：]?.*?(?:휴식|실시))"
        r"|(?:[·ㆍ-]?\s*W\s*/\s*L\s*카메라\s*[:：]?.*?(?:휴식|실시))",
        flags=re.IGNORECASE,
    )
    for measure in measures:
        # 기존 기록에 저장돼 있는 예전 카메라 운영 문구는 보고서에서 제외합니다.
        cleaned = legacy_camera_pattern.sub("", measure)
        cleaned = re.sub(r"\s*(?:\||/)+\s*(?:\||/)*\s*", " / ", cleaned)
        cleaned = cleaned.strip(" ·ㆍ-|/")
        if cleaned:
            cleaned_measures.append(cleaned)
    parts = [item.replace(" | ", " / ") for item in cleaned_measures]
    return "- " + (" / ".join(parts) if parts else "조치사항 기록 없음")


def report_department_measures(records: pd.DataFrame) -> str:
    """선택한 시행 조치와 직접 입력 조치를 부서별로 정리합니다."""
    lines: list[str] = []
    for department in TEAM_OPTIONS:
        aliases = {department, f"{department}팀"}
        if records[records["팀"].isin(aliases)].empty:
            continue
        summary = report_team_summary(
            records,
            f"{department}팀",
        ).removeprefix("- ")
        lines.append(f"  - {department} : {summary}")
    return "\n".join(lines) if lines else "  - 시행 조치 기록 없음"


def report_team_table_text(records: pd.DataFrame, department: str) -> str:
    """새 보고서 표에 들어갈 시행 조치를 정리합니다."""
    aliases = {department, f"{department}팀"}
    team_rows = records[records["팀"].isin(aliases)]
    if team_rows.empty:
        return "조치사항 기록 없음"

    measures: list[str] = []
    adjustment_reason = ""
    for value in team_rows["조치사항"].tolist():
        normalized_value = clean_text(value)
        # 일부 기존 기록은 괄호 안의 쉼표 위치가 조치 구분자(|)로
        # 저장되어 한 문장이 두 개의 글머리표로 출력됩니다.
        # 알려진 버튼 문구는 먼저 원래 한 문장으로 복원합니다.
        normalized_value = re.sub(
            r"휴식\s*공간\(휴게실\s*(?:,|\|)\s*그늘막\s*등\)\s*운영",
            "휴식 공간(휴게실, 그늘막 등) 운영",
            normalized_value,
        )
        normalized_value = re.sub(
            r"냉방\s*장치\s*가동\s*\(중계차\s*(?:,|\|)\s*방송\s*시설\s*등\)",
            "냉방 장치 가동 (중계차, 방송 시설 등)",
            normalized_value,
        )
        for part in re.split(r"\s*\|\s*", normalized_value):
            # 기존 시트 값에 수동 줄바꿈이 포함되어도 한 조치 문장은
            # 보고서에서 끊어지지 않도록 문장 내부 공백을 정규화합니다.
            item = re.sub(r"\s+", " ", part).strip()
            if not item:
                continue
            if re.search(
                r"필드\s*카메라|W\s*/\s*L\s*카메라",
                item,
                flags=re.IGNORECASE,
            ):
                continue
            if item.startswith(TIME_ADJUST_REASON_PREFIX):
                adjustment_reason = item[len(TIME_ADJUST_REASON_PREFIX):].strip()
                continue
            if item == "근무 시간대 조정":
                continue
            if item not in measures:
                measures.append(item)

    if adjustment_reason:
        measures.append(f"근무 시간대 조정 : {adjustment_reason}")
    return "\n".join(f"· {item}" for item in measures) or "조치사항 기록 없음"


def report_all_team_table_text(records: pd.DataFrame) -> str:
    """보고서 표 한 칸에 모든 부서의 시행 조치를 합칩니다."""
    lines: list[str] = []
    for department in TEAM_OPTIONS:
        aliases = {department, f"{department}팀"}
        if records[records["팀"].isin(aliases)].empty:
            continue
        label = f"{department}팀" if department != "기타" else "기타"
        measure_text = report_team_table_text(records, department)
        if measure_text == "조치사항 기록 없음":
            continue
        lines.append(f"{label}\n{measure_text}")
    return "\n\n".join(lines) or "조치사항 기록 없음"


def fill_report_measure_table(section_xml: str, records: pd.DataFrame) -> str:
    """새 원본의 3번 표를 부서별 행으로 나누지 않고 한 행에 채웁니다."""
    marker_index = section_xml.find("<hp:t>중계팀</hp:t>")
    if marker_index < 0:
        return section_xml
    table_start = section_xml.rfind("<hp:tbl", 0, marker_index)
    table_end = section_xml.find("</hp:tbl>", marker_index)
    if table_start < 0 or table_end < 0:
        return section_xml
    table_end += len("</hp:tbl>")
    table_xml = section_xml[table_start:table_end]
    row_match = re.search(r"<hp:tr>.*?</hp:tr>", table_xml, re.DOTALL)
    if not row_match:
        return section_xml

    template_row = row_match.group(0)
    rows: list[str] = []
    values = ["전체", report_all_team_table_text(records)]
    value_index = 0

    def replace_cell_text(match: re.Match[str]) -> str:
        nonlocal value_index
        value = values[value_index] if value_index < len(values) else ""
        value_index += 1
        return f"<hp:t>{html.escape(value, quote=False)}</hp:t>"

    row_xml = re.sub(
        r"<hp:t>.*?</hp:t>",
        replace_cell_text,
        template_row,
        flags=re.DOTALL,
    )
    row_xml = re.sub(r'rowAddr="\d+"', 'rowAddr="0"', row_xml)
    rows.append(row_xml)

    new_table = table_xml[:row_match.start()] + "".join(rows) + table_xml[row_match.end():]
    # 기존 예시 행은 모두 제거하고 한 행만 남깁니다.
    remaining_row = re.search(r"<hp:tr>.*?</hp:tr>", new_table[new_table.find(rows[-1]) + len(rows[-1]):], re.DOTALL)
    if remaining_row:
        offset = new_table.find(rows[-1]) + len(rows[-1])
        new_table = new_table[:offset + remaining_row.start()] + new_table[offset + remaining_row.end():]
    row_count = len(rows)
    new_table = re.sub(r'rowCnt="\d+"', f'rowCnt="{row_count}"', new_table, count=1)
    new_table = re.sub(
        r'(<hp:sz width="46850" widthRelTo="ABSOLUTE" height=")\d+("[^>]*/>)',
        rf"\g<1>{8074 * row_count}\g<2>",
        new_table,
        count=1,
    )
    return section_xml[:table_start] + new_table + section_xml[table_end:]


def report_team_work_time(records: pd.DataFrame, team: str) -> str:
    aliases = {team, team.removesuffix("팀")}
    team_rows = records[records["팀"].isin(aliases)]
    department = team.removesuffix("팀")
    if team_rows.empty:
        return f"{department} -"

    starts = sorted(unique_texts(team_rows["근무시작"].tolist()))
    ends = sorted(unique_texts(team_rows["근무종료"].tolist()))
    if not starts or not ends:
        return f"{department} -"
    return f"{department} {starts[0]}~{ends[-1]}"


def report_team_supervisor(records: pd.DataFrame, team: str) -> str:
    aliases = {team, team.removesuffix("팀")}
    team_rows = records[records["팀"].isin(aliases)]
    supervisors = unique_texts(team_rows["작성자"].tolist())
    department = team.removesuffix("팀")
    return f"{department} : {' / '.join(supervisors) if supervisors else '-'}"


def report_team_worker_count(records: pd.DataFrame, team: str) -> str:
    aliases = {team, team.removesuffix("팀")}
    team_rows = records[records["팀"].isin(aliases)]
    department = team.removesuffix("팀")
    if team_rows.empty:
        return f"{department} -"

    employees = sum(parse_int(value) for value in team_rows["직원"].tolist())
    contractors = sum(parse_int(value) for value in team_rows["도급"].tolist())
    saved_totals = sum(parse_int(value) for value in team_rows["근무자수"].tolist())
    total = saved_totals or (employees + contractors)
    return f"{department} 총 {total}명 (직원 {employees}명 / 도급 {contractors}명)"


def replace_report_text(xml_text: str, old: str, new: str) -> str:
    return xml_text.replace(old, html.escape(new, quote=False), 1)


def set_docx_cell_shading(cell: Any, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_docx_table_borders(table: Any, border_value: str = "nil") -> None:
    table_properties = table._tbl.tblPr
    borders = table_properties.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        table_properties.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = borders.find(qn(f"w:{edge}"))
        if border is None:
            border = OxmlElement(f"w:{edge}")
            borders.append(border)
        border.set(qn("w:val"), border_value)


def set_docx_cell_margins(cell: Any, top: int = 100, bottom: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for side, value in (("top", top), ("bottom", bottom), ("start", 120), ("end", 120)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def style_docx_run(run: Any, size: float = 10, bold: bool = False) -> None:
    run.font.name = "맑은 고딕"
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor(0, 0, 0)


def write_docx_cell(
    cell: Any,
    text: str,
    *,
    bold: bool = False,
    centered: bool = False,
    size: float = 9.5,
) -> None:
    cell.text = ""
    lines = text.splitlines() or [""]
    for index, line in enumerate(lines):
        paragraph = cell.paragraphs[0] if index == 0 else cell.add_paragraph()
        paragraph.alignment = (
            WD_ALIGN_PARAGRAPH.CENTER if centered else WD_ALIGN_PARAGRAPH.LEFT
        )
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15
        is_time_adjustment = line.lstrip("· ").startswith(
            "근무 시간대 조정 :"
        )
        style_docx_run(
            paragraph.add_run(line),
            size=size,
            bold=bold or is_time_adjustment,
        )
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_docx_cell_margins(cell)


def add_docx_section_heading(document: Any, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(9)
    paragraph.paragraph_format.space_after = Pt(5)
    paragraph.paragraph_format.keep_with_next = True
    style_docx_run(paragraph.add_run(text), size=11, bold=True)


def make_heat_report_docx_bytes(records: pd.DataFrame) -> bytes:
    """같은 날짜·현장의 기록을 줄 겹침 없는 Word 보고서로 생성합니다."""
    if records.empty:
        raise ValueError("보고서로 만들 기록이 없습니다.")

    first = records.iloc[0]
    work_date = clean_text(first.get("작업날짜"))
    try:
        work_date_text = datetime.strptime(work_date, "%Y-%m-%d").strftime("%Y.%m.%d")
    except ValueError:
        work_date_text = work_date.replace("-", ".")
    site = clean_text(first.get("현장명")) or "근무장소 미입력"
    sport = normalize_sport(first.get("종목"), site)

    heat_starts = sorted(unique_texts(records["폭염시작"].tolist()))
    heat_ends = sorted(unique_texts(records["폭염종료"].tolist()))
    heat_time = (
        f"{heat_starts[0]}~{heat_ends[-1]}"
        if heat_starts and heat_ends
        else "해당 없음"
    )
    temperatures: list[float] = []
    for value in records["체감온도"].tolist():
        temperatures.extend(temperature_numbers(value))
    if temperatures:
        low = format_number(min(temperatures))
        high = format_number(max(temperatures))
        temperature_text = f"{low}℃" if low == high else f"{low}℃~{high}℃"
    else:
        temperature_text = "미기록"

    departments = [
        department
        for department in TEAM_OPTIONS
        if not records[
            records["팀"].isin({department, f"{department}팀"})
        ].empty
    ]
    work_times = "\n".join(
        report_team_work_time(records, f"{department}팀")
        for department in departments
    ) or "-"
    authors = "\n".join(
        report_team_supervisor(records, f"{department}팀")
        for department in departments
    ) or "-"
    worker_counts = "\n".join(
        report_team_worker_count(records, f"{department}팀")
        for department in departments
    ) or "-"
    notes = " / ".join(unique_texts(records["특이사항"].tolist()))

    document = Document()
    section = document.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.45)
    section.bottom_margin = Cm(1.35)
    section.left_margin = Cm(1.55)
    section.right_margin = Cm(1.55)
    section.header_distance = Cm(0.8)
    section.footer_distance = Cm(0.8)

    normal = document.styles["Normal"]
    normal.font.name = "맑은 고딕"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "맑은 고딕")
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.15

    header_table = document.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    set_docx_table_borders(header_table, "nil")
    header_widths = [Inches(4.35), Inches(2.05)]
    for column, width in enumerate(header_widths):
        header_table.columns[column].width = width
        header_table.cell(0, column).width = width

    title_cell = header_table.cell(0, 0)
    title_cell.text = ""
    title_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    title = title_cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(2)
    style_docx_run(title.add_run("폭염작업 조치 결과 보고서"), size=16.5, bold=True)
    subtitle = title_cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(0)
    style_docx_run(
        subtitle.add_run("산업안전보건기준에 관한 규칙 제560조"),
        size=9,
    )

    approval_cell = header_table.cell(0, 1)
    approval_cell.text = ""
    approval_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    approval_table = approval_cell.add_table(rows=2, cols=3)
    approval_table.style = "Table Grid"
    approval_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    approval_table.autofit = False
    approval_widths = [Inches(0.38), Inches(0.82), Inches(0.85)]
    for column, width in enumerate(approval_widths):
        approval_table.columns[column].width = width
        for row in approval_table.rows:
            row.cells[column].width = width
    approval_label = approval_table.cell(0, 0).merge(approval_table.cell(1, 0))
    write_docx_cell(approval_label, "결\n재", bold=True, centered=True, size=9.5)
    write_docx_cell(approval_table.cell(0, 1), "담당", centered=True, size=9)
    write_docx_cell(approval_table.cell(0, 2), "센터장", centered=True, size=9)
    write_docx_cell(approval_table.cell(1, 1), "", centered=True, size=9)
    write_docx_cell(approval_table.cell(1, 2), "", centered=True, size=9)
    approval_table.rows[1].height = Cm(1.15)
    for paragraph in list(approval_cell.paragraphs):
        if not paragraph.text and paragraph._element.getnext() is approval_table._tbl:
            paragraph._element.getparent().remove(paragraph._element)

    after_header = document.add_paragraph()
    after_header.paragraph_format.space_after = Pt(1)

    add_docx_section_heading(document, "1. 기본 정보")
    info_table = document.add_table(rows=3, cols=4)
    info_table.style = "Table Grid"
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_table.autofit = False
    widths = [Inches(1.75), Inches(1.05), Inches(1.85), Inches(1.45)]
    headers = ["업무내용 / 장소", "근무일자", "근무시간", "작성자"]
    values = [f"{sport} / {site}", work_date_text, work_times, authors]
    for column, width in enumerate(widths):
        for row in info_table.rows:
            row.cells[column].width = width
        set_docx_cell_shading(info_table.cell(0, column), "E8EEF5")
        write_docx_cell(info_table.cell(0, column), headers[column], bold=True, centered=True)
        write_docx_cell(info_table.cell(1, column), values[column], centered=True)
    worker_value_cell = info_table.cell(2, 1).merge(info_table.cell(2, 3))
    set_docx_cell_shading(info_table.cell(2, 0), "E8EEF5")
    write_docx_cell(info_table.cell(2, 0), "근무인원", bold=True, centered=True)
    write_docx_cell(worker_value_cell, worker_counts, centered=True)

    add_docx_section_heading(document, "2. 폭염작업 현황")
    heat_table = document.add_table(rows=2, cols=2)
    heat_table.style = "Table Grid"
    heat_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    heat_table.autofit = False
    for column, (header, value) in enumerate(
        (("폭염시간대", heat_time), ("체감온도", temperature_text))
    ):
        for row in heat_table.rows:
            row.cells[column].width = Inches(3.05)
        set_docx_cell_shading(heat_table.cell(0, column), "E8EEF5")
        write_docx_cell(heat_table.cell(0, column), header, bold=True, centered=True)
        write_docx_cell(heat_table.cell(1, column), value, centered=True)
    source_note = document.add_paragraph()
    source_note.paragraph_format.space_before = Pt(3)
    source_note.paragraph_format.space_after = Pt(2)
    style_docx_run(
        source_note.add_run("※ 기상청 날씨누리 시간대별 예보 기준"),
        size=8.5,
    )

    add_docx_section_heading(document, "3. 폭염 대응 조치")
    measure_table = document.add_table(rows=1, cols=2)
    measure_table.style = "Table Grid"
    measure_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    measure_table.autofit = False
    measure_table.rows[0].cells[0].width = Inches(1.15)
    measure_table.rows[0].cells[1].width = Inches(4.95)
    set_docx_cell_shading(measure_table.cell(0, 0), "E8EEF5")
    write_docx_cell(
        measure_table.cell(0, 0),
        "전체",
        bold=True,
        centered=True,
    )
    write_docx_cell(
        measure_table.cell(0, 1),
        report_all_team_table_text(records),
        size=9.3,
    )

    if notes:
        add_docx_section_heading(document, "4. 특이사항")
        notes_paragraph = document.add_paragraph()
        notes_paragraph.paragraph_format.left_indent = Cm(0.35)
        notes_paragraph.paragraph_format.space_after = Pt(4)
        style_docx_run(notes_paragraph.add_run(notes), size=9.5)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_docx_run(
        footer.add_run("㈜후니드 TM사업부문 미디어제작센터"),
        size=8.5,
    )

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()


@st.cache_data(show_spinner=False, ttl=86400)
def cached_heat_report_docx_bytes(
    records_json: str,
    report_layout_version: str,
) -> bytes:
    """같은 기록의 Word 보고서를 24시간 캐시해 재생성을 줄입니다."""
    _ = report_layout_version
    records = pd.read_json(io.StringIO(records_json), orient="records")
    return make_heat_report_docx_bytes(records)


def report_attachment_for_rows(records: pd.DataFrame) -> tuple[str, bytes]:
    if records.empty:
        raise ValueError("보고서로 만들 기록이 없습니다.")
    records = canonicalize_report_rows(records)
    work_date = clean_text(records.iloc[0].get("작업날짜"))
    site = clean_text(records.iloc[0].get("현장명"))
    safe_site = re.sub(r"[^0-9A-Za-z가-힣_-]+", "_", site).strip("_") or "현장"
    filename = (
        "폭염작업_조치_결과_보고서_"
        f"{work_date.replace('-', '')}_{safe_site}.docx"
    )
    records_json = records.to_json(
        orient="records",
        force_ascii=False,
        date_format="iso",
    )
    report_bytes = cached_heat_report_docx_bytes(
        records_json,
        "approval-box-v7-sport-region-site-merge",
    )
    return filename, report_bytes


def find_nested_secret(container: Any, names: set[str]) -> str:
    """TOML의 어느 섹션에 들어가도 지정된 메일 키를 찾아 반환합니다."""
    if not isinstance(container, Mapping):
        return ""
    for key, value in container.items():
        if str(key).lower() in names and not isinstance(value, Mapping):
            return clean_text(value)
    for value in container.values():
        found = find_nested_secret(value, names)
        if found:
            return found
    return ""


def mail_secret(name: str, section_name: str) -> str:
    direct = get_secret((name,)) or get_secret(("mail", section_name))
    if direct:
        return direct
    try:
        secrets_tree: Any = st.secrets.to_dict()
    except Exception:  # noqa: BLE001
        secrets_tree = st.secrets
    return find_nested_secret(
        secrets_tree,
        {name.lower(), section_name.lower()},
    )


def report_mail_recipients() -> list[str]:
    raw = mail_secret("MAIL_RECIPIENTS", "recipients")
    return [
        item.strip()
        for item in re.split(r"[,;\n]+", raw)
        if item.strip()
    ]


def send_report_attachments_email(
    report_date: str,
    scope: str,
    attachments: list[tuple[str, bytes]],
) -> int:
    sender = mail_secret("MAIL_SENDER", "sender")
    password = mail_secret("MAIL_APP_PASSWORD", "app_password").replace(" ", "")
    recipients = report_mail_recipients()
    if not sender or not password or not recipients:
        missing = []
        if not sender:
            missing.append("MAIL_SENDER")
        if not password:
            missing.append("MAIL_APP_PASSWORD")
        if not recipients:
            missing.append("MAIL_RECIPIENTS")
        raise RuntimeError(
            "Streamlit Secrets에서 다음 항목을 찾지 못했습니다: "
            + ", ".join(missing)
        )

    if not attachments:
        raise RuntimeError("메일에 첨부할 보고서가 없습니다.")
    if sum(len(payload) for _, payload in attachments) > 20 * 1024 * 1024:
        raise RuntimeError("첨부파일 합계가 20MB를 넘어 날짜를 나누어 발송해 주세요.")

    with smtplib.SMTP_SSL("smtp.gmail.com", 465, timeout=30) as smtp:
        smtp.login(sender, password)
        for recipient in recipients:
            message = EmailMessage()
            message["Subject"] = (
                f"[CheckTemp] {report_date} {scope} 폭염작업 조치 결과 보고서"
            )
            message["From"] = sender
            message["To"] = recipient
            message.set_content(
                f"{report_date} {scope} 근무 기록으로 작성한 보고서입니다.\n\n"
                f"첨부된 Word 보고서 {len(attachments)}개를 확인해 주세요.\n\n"
                "이 메일은 CheckTemp에서 발송되었습니다."
            )
            for filename, report_bytes in attachments:
                message.add_attachment(
                    report_bytes,
                    maintype="application",
                    subtype=(
                        "vnd.openxmlformats-officedocument."
                        "wordprocessingml.document"
                    ),
                    filename=filename,
                )
            smtp.send_message(message)
    return len(recipients)


def send_report_email(
    report_date: str,
    site: str,
    filename: str,
    report_bytes: bytes,
) -> int:
    return send_report_attachments_email(
        report_date,
        site,
        [(filename, report_bytes)],
    )


def make_heat_report_bytes(records: pd.DataFrame) -> bytes:
    """같은 날짜·현장의 기록을 원본 HWPX 양식에 채웁니다."""
    if records.empty:
        raise ValueError("보고서로 만들 기록이 없습니다.")

    template_path = Path(__file__).with_name("report_template.b64")
    template_bytes = base64.b64decode(
        template_path.read_text(encoding="ascii").strip()
    )

    first = records.iloc[0]
    work_date = clean_text(first.get("작업날짜"))
    try:
        work_date_text = datetime.strptime(
            work_date,
            "%Y-%m-%d",
        ).strftime("%Y.%m.%d")
    except ValueError:
        work_date_text = work_date.replace("-", ".")

    site = clean_text(first.get("현장명")) or "현장명 미입력"
    sport = normalize_sport(first.get("종목"), site)
    site_text = f"{sport} / {site}"

    heat_starts = sorted(unique_texts(records["폭염시작"].tolist()))
    heat_ends = sorted(unique_texts(records["폭염종료"].tolist()))
    heat_time = (
        f"{heat_starts[0]}~{heat_ends[-1]}"
        if heat_starts and heat_ends
        else "해당 없음"
    )

    temperatures: list[float] = []
    for value in records["체감온도"].tolist():
        temperatures.extend(temperature_numbers(value))
    if temperatures:
        low = format_number(min(temperatures))
        high = format_number(max(temperatures))
        temperature_text = (
            f"{low}℃" if low == high else f"{low}℃~{high}℃"
        )
    else:
        temperature_text = "미기록"

    common_values = unique_texts(records["공통 조치사항"].tolist())
    common_text = common_values[0] if common_values else common_measures_for_sport(sport)
    common_lines = [
        line.strip().lstrip("- ").strip()
        for line in common_text.splitlines()
        if line.strip()
    ]
    while len(common_lines) < 3:
        common_lines.append("")

    notes = " / ".join(unique_texts(records["특이사항"].tolist()))

    replacements = [
        ("기타/ENG(실내) / SBS프리즘타워", site_text),
        ("야구 / 삼성야구장", site_text),
        ("종목 / 장소명", site_text),
        ("2026.08.12", work_date_text),
        ("2026.08.13", work_date_text),
        ("2026.08.10", work_date_text),
        ("중계팀 18:30~21:30", report_team_work_time(records, "중계팀")),
        ("영상팀 18:30~21:30", report_team_work_time(records, "영상팀")),
        ("중계팀 08:00~17:00", report_team_work_time(records, "중계팀")),
        ("영상팀 08:00~17:00", report_team_work_time(records, "영상팀")),
        ("중계팀 : 홍길동", report_team_supervisor(records, "중계팀")),
        ("영상팀 : 홍길동", report_team_supervisor(records, "영상팀")),
        ("중계 : 박준상2", report_team_supervisor(records, "중계팀")),
        ("영상 : 박준상", report_team_supervisor(records, "영상팀")),
        ("중계 9:00~13:00", report_team_work_time(records, "중계팀")),
        ("영상 9:00~13:00", report_team_work_time(records, "영상팀")),
        ("해당 없음", heat_time),
        ("08:00~15:00", heat_time),
        ("28.9℃~30.9℃", temperature_text),
        ("31℃~33℃", temperature_text),
        ("12:00~13:00", heat_time),
        ("29.6℃", temperature_text),
        (
            "- 중계차·장비차·중계석·중계스태프실 냉방 가동",
            f"- {common_lines[0]}",
        ),
        (
            "냉방 공간 체감온도 27℃ 이하 유지",
            common_lines[1],
        ),
        (
            "- 생수·냉음료·식염포도당·폭염질환 응급키트 비치 및 위치 공유",
            f"- {common_lines[2]}",
        ),
        (
            "- 중계차, 중계석, 휴게실 냉방 가동(체감온도 OO℃ 이하 유지)",
            f"- {common_lines[0]}",
        ),
        (
            "식염포도당, 폭염질환 응급키트 위치 공유 및 생수 및 이온 음료",
            common_lines[1],
        ),
        (" 지급", ""),
        ("  - 폭염 시간대 불필요한 외부 활동 최소화", f"  - {common_lines[2]}"),
        (
            "  - 예시 1) 제작팀 협의 사항 적용 : ",
            "  - " if notes else "",
        ),
        ("강조효과 테스트 2", ""),
        ("영상팀 필드카메라 중요 선수 외 이글 또는 버디", ""),
        ("상황까지", ""),
        ("만", ""),
        ("                                     촬영하고 그 외 추가 휴식시간 부여", ""),
        ("  - 예시 2) 일정 조정 : 기존 출근시간 대비 1시간 조기 출근하여 실외 작업을 조기", ""),
        ("                       진행하고, 폭염시간대에는 1시간 이내 20분 이상 휴게시간을", ""),
        ("                       확보하여 운영", ""),
        ("- 중계차 및 중계석 등 냉방공간 근무 → ", "__CHECKTEMP_BROADCAST_MEASURES__"),
        ("- 조치사항 기록 없음", "__CHECKTEMP_BROADCAST_MEASURES__"),
        ("폭염작업 해당 없음", ""),
        ("- 1시간 이내 10분 이상 휴게시간 부여", "__CHECKTEMP_VIDEO_MEASURES__"),
        (
            "- 1시간 이내 10분 이상 휴식 / 누적 휴게시간 10분",
            "__CHECKTEMP_VIDEO_MEASURES__",
        ),
        (" · 필드카메라 : 중요 선수 외 이글·버디 워킹 촬영 후 휴식", ""),
        (" · W/L 카메라 : 혹서기 지원인력을 활용한 교대근무 실시", ""),
    ]

    source = io.BytesIO(template_bytes)
    output = io.BytesIO()
    with zipfile.ZipFile(source, "r") as input_zip:
        with zipfile.ZipFile(output, "w") as output_zip:
            for info in input_zip.infolist():
                payload = input_zip.read(info.filename)
                if info.filename == "Contents/section0.xml":
                    section_xml = payload.decode("utf-8")
                    for old, new in replacements:
                        section_xml = replace_report_text(section_xml, old, new)

                    # 사용자가 수정한 최신 원본의 팀별 표를 실제 기록으로 채웁니다.
                    section_xml = fill_report_measure_table(
                        section_xml,
                        records,
                    )

                    # 치환된 조치 문구가 다음 검색어로 다시 오인되지 않도록
                    # 마지막 단계에서 팀별 임시표시를 실제 내용으로 바꿉니다.
                    section_xml = section_xml.replace(
                        "__CHECKTEMP_BROADCAST_MEASURES__",
                        html.escape(
                            report_team_summary(records, "중계팀"),
                            quote=False,
                        ),
                    )
                    section_xml = section_xml.replace(
                        "__CHECKTEMP_VIDEO_MEASURES__",
                        html.escape(
                            report_team_summary(records, "영상팀"),
                            quote=False,
                        ),
                    )

                    # 예전 원본에서 휴게 조치 결과 칸에 남아 있던
                    # 필드카메라/W/L 카메라 안내문은 문장 분할·공백과 관계없이 제거합니다.
                    report_result_marker = "4. 휴게 조치 결과"
                    marker_index = section_xml.find(report_result_marker)
                    if marker_index >= 0:
                        report_head = section_xml[:marker_index]
                        report_result = section_xml[marker_index:]
                        legacy_camera_patterns = (
                            r"\s*[·ㆍ-]?\s*필드카메라\s*[:：]?\s*"
                            r".*?(?:휴식|실시)\s*",
                            r"\s*[·ㆍ-]?\s*W\s*/\s*L\s*카메라\s*[:：]?\s*"
                            r".*?(?:휴식|실시)\s*",
                        )

                        def remove_legacy_camera_text(match: re.Match[str]) -> str:
                            value = html.unescape(match.group(1))
                            for pattern in legacy_camera_patterns:
                                value = re.sub(pattern, "", value)
                            return f"<hp:t>{html.escape(value, quote=False)}</hp:t>"

                        report_result = re.sub(
                            r"<hp:t>(.*?)</hp:t>",
                            remove_legacy_camera_text,
                            report_result,
                            flags=re.DOTALL,
                        )
                        section_xml = report_head + report_result

                    # 선택 버튼과 '그 외 조치사항'을 폭염 대응 조치 안에
                    # 중계·영상·기타 부서별로 표시합니다.
                    schedule_heading = (
                        "<hp:t> 2) 일정 조정 또는 제작팀 협의 사항"
                        "(없는 경우 생략)</hp:t>"
                    )
                    department_measures = html.escape(
                        report_department_measures(records),
                        quote=False,
                    )
                    schedule_replacement = (
                        "<hp:t> 2) 부서별 시행 조치</hp:t>"
                        "<hp:lineBreak/>"
                        f"<hp:t>{department_measures}</hp:t>"
                    )
                    section_xml = section_xml.replace(
                        schedule_heading,
                        schedule_replacement,
                        1,
                    )

                    # 한글 보고서에서는 공통 예방조치 제목과 내용을 제외합니다.
                    # 앱과 스프레드시트에 저장된 공통 조치 데이터는 유지됩니다.
                    common_heading_index = section_xml.find(
                        "<hp:t> 1) 공통 예방조치</hp:t>"
                    )
                    department_heading_index = section_xml.find(
                        "<hp:t> 2) 부서별 시행 조치</hp:t>"
                    )
                    if (
                        common_heading_index >= 0
                        and department_heading_index > common_heading_index
                    ):
                        common_paragraph_start = section_xml.rfind(
                            "<hp:p",
                            0,
                            common_heading_index,
                        )
                        department_paragraph_start = section_xml.rfind(
                            "<hp:p",
                            0,
                            department_heading_index,
                        )
                        if (
                            common_paragraph_start >= 0
                            and department_paragraph_start
                            > common_paragraph_start
                        ):
                            section_xml = (
                                section_xml[:common_paragraph_start]
                                + section_xml[department_paragraph_start:]
                            )

                    # 기존 '4. 휴게 조치 결과' 표는 제거합니다. 앱의 특이사항은
                    # 입력값이 있을 때만 4번 항목으로 작성합니다.
                    result_marker = "4. 휴게 조치 결과"
                    result_index = section_xml.find(result_marker)
                    if result_index >= 0:
                        table_start = section_xml.find("<hp:tbl", result_index)
                        table_end = section_xml.find("</hp:tbl>", table_start)
                        if table_start >= 0 and table_end >= 0:
                            table_end += len("</hp:tbl>")
                            section_xml = (
                                section_xml[:table_start]
                                + section_xml[table_end:]
                            )
                        consultation_text = (
                            "4. 특이사항\n  - "
                            + html.escape(notes, quote=False)
                            if notes
                            else ""
                        )
                        section_xml = section_xml.replace(
                            result_marker,
                            consultation_text,
                            1,
                        )

                    # 최신 원본의 4번 항목은 제목과 내용을 반드시 줄바꿈합니다.
                    # 특이사항이 없으면 해당 문단 전체를 생략합니다.
                    special_paragraph_pattern = (
                        r"<hp:p(?P<attrs>[^>]*)>"
                        r"(?P<body>.*?<hp:t>4\. 특이사항.*?</hp:t>.*?)"
                        r"</hp:p>"
                    )
                    special_match = re.search(
                        special_paragraph_pattern,
                        section_xml,
                        flags=re.DOTALL,
                    )
                    if special_match:
                        if notes:
                            body = special_match.group("body")
                            body = re.sub(
                                r"<hp:t>4\. 특이사항.*?</hp:t>",
                                "<hp:t>4. 특이사항<hp:lineBreak/>"
                                f"  - {html.escape(notes, quote=False)}</hp:t>",
                                body,
                                count=1,
                                flags=re.DOTALL,
                            )
                            special_replacement = (
                                f"<hp:p{special_match.group('attrs')}>"
                                f"{body}</hp:p>"
                            )
                        else:
                            special_replacement = ""
                        section_xml = (
                            section_xml[:special_match.start()]
                            + special_replacement
                            + section_xml[special_match.end():]
                        )
                    # 새 원본에 예시로 입력된 특이사항 문단은 실제 값과 별개이므로 제거합니다.
                    section_xml = re.sub(
                        r"<hp:p[^>]*>(?:(?!</hp:p>).)*설치 철수 등"
                        r"(?:(?!</hp:p>).)*</hp:p>",
                        "",
                        section_xml,
                        count=1,
                        flags=re.DOTALL,
                    )
                    payload = section_xml.encode("utf-8")
                output_zip.writestr(info, payload)

    return output.getvalue()


def make_excel_bytes(dataframe: pd.DataFrame) -> bytes:
    """전체 기록을 관리용 XLSX 파일로 생성합니다."""
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "전체기록"

    header_fill = PatternFill("solid", fgColor="172B4D")
    header_font = Font(color="FFFFFF", bold=True)
    zebra_fill = PatternFill("solid", fgColor="F6F8FB")
    notes_fill = PatternFill("solid", fgColor="FFECCD")
    notes_font = Font(color="8C330D", bold=True)

    thin_side = Side(style="thin", color="C4CAD4")
    grid_border = Border(
        left=thin_side,
        right=thin_side,
        top=thin_side,
        bottom=thin_side,
    )

    for column_index, column_name in enumerate(COLUMNS, start=1):
        display_name = {
            "작업날짜": "근무일자",
            "종목": "업무내용",
            "현장명": "근무장소",
            "팀": "부서",
        }.get(column_name, column_name)
        cell = worksheet.cell(row=1, column=column_index, value=display_name)
        cell.fill = header_fill
        cell.font = header_font
        cell.border = grid_border
        cell.alignment = Alignment(
            horizontal="center",
            vertical="center",
            wrap_text=True,
        )

    export_df = dataframe.reindex(columns=COLUMNS).fillna("")

    for row_index, (_, row) in enumerate(export_df.iterrows(), start=2):
        for column_index, column_name in enumerate(COLUMNS, start=1):
            value = clean_text(row.get(column_name, ""))
            cell = worksheet.cell(
                row=row_index,
                column=column_index,
                value=value,
            )
            cell.border = grid_border
            cell.alignment = Alignment(
                vertical="center",
                wrap_text=True,
            )
            if row_index % 2 == 1:
                cell.fill = zebra_fill

        notes_cell = worksheet.cell(
            row=row_index,
            column=COLUMNS.index("특이사항") + 1,
        )
        if clean_text(notes_cell.value):
            notes_cell.fill = notes_fill
            notes_cell.font = notes_font

    worksheet.freeze_panes = "A2"
    worksheet.auto_filter.ref = (
        f"A1:{get_column_letter(len(COLUMNS))}{max(1, worksheet.max_row)}"
    )

    preferred_widths = {
        "id": 18,
        "작업날짜": 12,
        "종목": 12,
        "현장명": 24,
        "팀": 10,
        "근무자수": 10,
        "직원": 8,
        "도급": 8,
        "근무시작": 10,
        "근무종료": 10,
        "작성자": 12,
        "폭염시작": 10,
        "폭염종료": 10,
        "체감온도": 12,
        "휴게시간": 10,
        "공통 조치사항": 48,
        "조치사항": 34,
        "특이사항": 36,
        "등록시간": 20,
        "수정시간": 20,
    }

    for index, column_name in enumerate(COLUMNS, start=1):
        worksheet.column_dimensions[
            get_column_letter(index)
        ].width = preferred_widths.get(column_name, 16)

    worksheet.row_dimensions[1].height = 26

    output = io.BytesIO()
    workbook.save(output)
    return output.getvalue()


def init_state() -> None:
    defaults = {
        "page": "form",
        "editing_id": None,
        "pending_delete": None,
        "form_nonce": 0,
        "flash": "",
    }

    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


def reset_form(*, go_to_form: bool = True) -> None:
    st.session_state.editing_id = None
    st.session_state.pending_delete = None
    st.session_state.form_nonce += 1

    if go_to_form:
        st.session_state.page = "form"


def show_flash() -> None:
    message = clean_text(
        st.session_state.get("flash", "")
    )

    if message:
        st.success(message)
        st.session_state.flash = ""


def render_section_heading(
    number: str,
    title: str,
    subtitle: str,
) -> None:
    st.markdown(
        f"""
        <div class="section-heading">
            <span class="section-number">
                {html.escape(number)}
            </span>
            <div>
                <b>{html.escape(title)}</b>
                <small>{html.escape(subtitle)}</small>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_header() -> None:
    st.markdown(
        f"""
        <div class="app-header">
            <div>
                <div class="app-kicker">현장 안전관리</div>
                <h1 class="app-title">
                    폭염대비 온열질환 예방을 위한 조치사항
                </h1>
                <p class="app-subtitle">
                    폭염작업이란 31℃ 이상인 작업장소에서
                    연속 2시간 이상 작업하는 것을 말합니다.
                </p>
            </div>
            <div class="app-version">
                {html.escape(APP_VERSION)}
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_navigation() -> None:
    left, right = st.columns(2)

    with left:
        if st.button(
            "새 기록",
            use_container_width=True,
            type=(
                "primary"
                if st.session_state.page == "form"
                else "secondary"
            ),
            key="nav_form",
        ):
            reset_form(go_to_form=True)
            st.rerun()

    with right:
        if st.button(
            "기록 조회",
            use_container_width=True,
            type=(
                "primary"
                if st.session_state.page == "records"
                else "secondary"
            ),
            key="nav_records",
        ):
            st.session_state.page = "records"
            st.session_state.pending_delete = None
            st.rerun()


def render_setup_error(error: Exception) -> None:
    st.markdown(
        """
        <div class="setup-box">
        <b>Google Sheets 연결 설정이 아직 완료되지 않았습니다.</b>
        <br>
        화면은 미리 볼 수 있지만, 연결 전에는
        저장·수정·삭제가 비활성화됩니다.
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.expander("설정 오류 자세히 보기"):
        st.code(str(error))
        st.caption(
            "Streamlit App settings → Secrets에 "
            "스프레드시트 주소와 서비스 계정 정보를 "
            "등록해야 합니다."
        )


def render_form(
    records: pd.DataFrame,
    store_error: Exception | None,
) -> None:
    editing_id = clean_text(
        st.session_state.editing_id
    )
    editing_record: dict[str, Any] = {}

    if editing_id and not records.empty:
        matches = records[
            records["id"].astype(str) == editing_id
        ]

        if not matches.empty:
            editing_record = matches.iloc[0].to_dict()
        else:
            st.warning(
                "수정하려던 기록을 찾지 못해 "
                "새 기록 화면으로 전환했습니다."
            )
            reset_form(go_to_form=True)
            st.rerun()

    title = "기록 수정" if editing_id else "새 기록 작성"
    st.markdown(f"## {title}")

    nonce = st.session_state.form_nonce
    initialize_time_state(editing_record, nonce)
    initialize_rest_minutes(editing_record, nonce)
    initialize_team_state(editing_record, nonce)
    initialize_measures_state(editing_record, nonce)

    default_date = datetime.now(KST).date()
    date_text = clean_text(
        editing_record.get("작업날짜")
    )

    if date_text:
        try:
            default_date = date.fromisoformat(date_text)
        except ValueError:
            pass

    temperature_default = clean_text(
        editing_record.get("체감온도")
    )

    with st.container(border=True):
        render_section_heading(
            "01",
            "기본 정보",
            "현장과 담당 정보를 입력합니다.",
        )

        work_date = st.date_input(
            "근무일자 *",
            value=default_date,
            key=f"date_{nonce}",
            on_change=auto_lookup_future_weather,
            args=(nonce,),
        )

        sport = st.selectbox(
            "업무내용 *",
            SPORT_OPTIONS,
            index=option_index(
                SPORT_OPTIONS,
                normalize_sport(
                    editing_record.get("종목"),
                    editing_record.get("현장명"),
                ),
            ),
            key=f"sport_{nonce}",
        )

        site = st.text_input(
            "근무장소 *",
            value=clean_text(
                editing_record.get("현장명")
            ),
            placeholder="예: ○○골프장, ○○야구장",
            key=f"site_{nonce}",
            on_change=auto_lookup_future_weather,
            args=(nonce,),
        )

        st.markdown(
            '<div class="field-label">부서 선택 *</div>',
            unsafe_allow_html=True,
        )

        selected_team = clean_text(
            st.session_state.get(team_state_key(nonce))
        )
        team_left, team_center, team_right = st.columns(3)

        with team_left:
            st.button(
                "중계",
                key=f"team_relay_{nonce}",
                use_container_width=True,
                type=(
                    "primary"
                    if selected_team == "중계"
                    else "secondary"
                ),
                on_click=set_team,
                args=("중계", nonce),
            )

        with team_center:
            st.button(
                "영상",
                key=f"team_video_{nonce}",
                use_container_width=True,
                type=(
                    "primary"
                    if selected_team == "영상"
                    else "secondary"
                ),
                on_click=set_team,
                args=("영상", nonce),
            )

        with team_right:
            st.button(
                "기타",
                key=f"team_other_{nonce}",
                use_container_width=True,
                type=(
                    "primary"
                    if selected_team == "기타"
                    else "secondary"
                ),
                on_click=set_team,
                args=("기타", nonce),
            )

        team = clean_text(
            st.session_state.get(team_state_key(nonce))
        )

        worker_count = 0
        employee_count = 0
        contractor_count = 0
        work_start_input = ""
        work_end_input = ""

        if team:
            total_col, employee_col, contractor_col = st.columns(3)
            headcount_options = list(range(0, 501))
            worker_count_key = f"worker_count_{nonce}"
            if worker_count_key not in st.session_state:
                saved_worker_count = parse_int(
                    editing_record.get("근무자수"),
                    -1,
                )
                if saved_worker_count >= 0:
                    st.session_state[worker_count_key] = saved_worker_count
                else:
                    st.session_state[worker_count_key] = (
                        parse_int(editing_record.get("직원"))
                        + parse_int(editing_record.get("도급"))
                    )
            with employee_col:
                employee_count = st.selectbox(
                    "직원",
                    headcount_options,
                    index=min(parse_int(editing_record.get("직원")), 500),
                    key=f"employee_count_{nonce}",
                    format_func=lambda value: f"{value}명",
                    on_change=sync_worker_count,
                    args=(nonce,),
                )
            with contractor_col:
                contractor_count = st.selectbox(
                    "도급",
                    headcount_options,
                    index=min(parse_int(editing_record.get("도급")), 500),
                    key=f"contractor_count_{nonce}",
                    format_func=lambda value: f"{value}명",
                    on_change=sync_worker_count,
                    args=(nonce,),
                )
            with total_col:
                worker_count = st.number_input(
                    "근무자 수",
                    min_value=0,
                    max_value=1000,
                    step=1,
                    key=worker_count_key,
                )

            st.caption(
                "직원과 도급 인원을 선택하면 자동 합산됩니다. "
                "필요하면 근무자 수를 직접 수정할 수도 있습니다."
            )

            st.caption(
                "근무시간은 24시간 형식으로 입력하세요. "
                "예: 오전 9시는 09:00, 오후 6시는 18:00"
            )

            work_left, work_right = st.columns(2)

            with work_left:
                work_start_input = st.text_input(
                    "근무 시작 *",
                    value=clean_text(editing_record.get("근무시작")),
                    placeholder="예: 2256 → 22:56",
                    key=f"manual_work_start_{nonce}",
                    on_change=normalize_manual_time_field,
                    args=("work_start", nonce),
                    kwargs={"trigger_weather": True},
                )

            with work_right:
                work_end_input = st.text_input(
                    "근무 종료 *",
                    value=clean_text(editing_record.get("근무종료")),
                    placeholder="예: 1830 → 18:30",
                    key=f"manual_work_end_{nonce}",
                    on_change=normalize_manual_time_field,
                    args=("work_end", nonce),
                    kwargs={"trigger_weather": True},
                )

            st.caption(
                "숫자만 입력해도 됩니다. Enter를 누르면 "
                "2256→22:56, 930→09:30, 18→18:00으로 자동 변환됩니다."
            )
        else:
            st.caption("부서를 선택하면 근무 인원과 근무시간 입력란이 표시됩니다.")

        author = st.text_input(
            "작성자",
            value=clean_text(
                editing_record.get("작성자")
            ),
            placeholder="예: 홍길동",
            key=f"author_{nonce}",
        )

        st.divider()

        render_section_heading(
            "02",
            "폭염 정보",
            "폭염작업 시간과 체감온도를 기록합니다.",
        )

        st.caption(
            "폭염시간은 24시간 형식으로 입력하세요. "
            "예: 오후 1시 30분은 13:30"
        )

        heat_left, heat_right = st.columns(2)

        with heat_left:
            heat_start_input = st.text_input(
                "폭염작업 시작",
                value=clean_text(editing_record.get("폭염시작")),
                placeholder="예: 1330 → 13:30",
                key=f"manual_heat_start_{nonce}",
                on_change=normalize_manual_time_field,
                args=("heat_start", nonce),
            )

        with heat_right:
            heat_end_input = st.text_input(
                "폭염작업 종료",
                value=clean_text(editing_record.get("폭염종료")),
                placeholder="예: 1700 → 17:00",
                key=f"manual_heat_end_{nonce}",
                on_change=normalize_manual_time_field,
                args=("heat_end", nonce),
            )

        work_start = parse_time_value(work_start_input)
        work_end = parse_time_value(work_end_input)
        heat_start = parse_time_value(heat_start_input)
        heat_end = parse_time_value(heat_end_input)

        st.button(
            "체감온도 자동 조회",
            key=f"lookup_weather_{nonce}",
            use_container_width=True,
            on_click=record_heat_start_with_weather,
            args=(nonce, False),
        )

        cache_restored = restore_weather_result_cache(nonce)
        if cache_restored:
            st.caption("동일한 조회 조건의 최근 결과를 캐시에서 불러왔습니다.")

        weather_notice = st.session_state.get(
            weather_notice_key(nonce)
        )
        if weather_notice:
            notice_type, notice_message = weather_notice
            if notice_type == "success":
                st.success(notice_message)
            elif notice_type == "warning":
                st.warning(notice_message)
            else:
                st.info(notice_message)

        weather_debug_lines = st.session_state.get(
            weather_debug_log_key(nonce),
            [],
        )
        if weather_debug_lines:
            with st.expander("기상조회 진단 로그", expanded=False):
                st.caption(
                    "오류 분석용 로그입니다. 인증키는 기록되지 않습니다. "
                    "조회 실패 시 아래 내용을 복사해 전달해 주세요."
                )
                st.code(
                    "\n".join(weather_debug_lines),
                    language="text",
                )

        is_forecast_entry = work_date > datetime.now(KST).date()
        if is_forecast_entry:
            st.info(
                "미래 날짜 사전입력입니다. 근무장소와 근무시간을 모두 "
                "입력하면 단기예보 체감온도와 예상 폭염시간이 자동 "
                "적용됩니다. 작업 후 기록 조회의 수정 버튼으로 실제값을 "
                "반영할 수 있습니다."
            )

        temperature = st.text_input(
            "체감온도 (℃)",
            value=temperature_default,
            placeholder="예: 33~35 또는 34",
            help=(
                "단일값은 34, 범위는 33~35처럼 "
                "입력합니다. ℃는 입력하지 않아도 됩니다."
            ),
            key=f"temperature_{nonce}",
        )

        st.caption(
            "미래 날짜는 근무시간의 단기예보를 이용해 예상 체감온도와 "
            "31℃ 이상 예상 폭염시간을 자동 입력합니다. 오늘 또는 지난 "
            "날짜는 기존처럼 폭염시간을 우선 조회하고, 비어 있으면 "
            "근무시간의 기상청 500m 격자값을 조회합니다."
        )

        st.divider()

        render_section_heading(
            "03",
            "조치 사항",
            "시행한 조치와 특이사항을 남깁니다.",
        )

        common_measures = ""

        st.markdown(
            '<div class="field-label">시행 조치</div>',
            unsafe_allow_html=True,
        )
        st.caption("복수 선택 가능합니다. 다시 누르면 선택이 해제됩니다.")

        selected_measure_list = list(
            st.session_state.get(
                measures_state_key(nonce),
                [],
            )
        )

        measure_columns = st.columns(2)
        for index, measure in enumerate(MEASURE_OPTIONS):
            with measure_columns[index % 2]:
                st.button(
                    measure,
                    key=f"measure_{index}_{nonce}",
                    use_container_width=True,
                    type=(
                        "primary"
                        if measure in selected_measure_list
                        else "secondary"
                    ),
                    on_click=toggle_measure,
                    args=(measure, nonce),
                )

        measures = list(
            st.session_state.get(
                measures_state_key(nonce),
                [],
            )
        )

        time_adjust_reason = ""
        if "근무 시간대 조정" in measures:
            time_adjust_reason = st.text_input(
                "근무시간 조정 사유 *",
                value=time_adjust_reason_text(
                    editing_record.get("조치사항")
                ),
                placeholder="예: 폭염시간대를 피해 설치 시간을 앞당김",
                key=f"time_adjust_reason_{nonce}",
            )

        other_measure = st.text_area(
            "그 외 조치사항",
            value=custom_measure_text(
                editing_record.get("조치사항")
            ),
            placeholder="예: 냉음료 추가 지급, 작업 순서 조정",
            height=90,
            key=f"other_measure_{nonce}",
        )

        notes = st.text_area(
            "특이사항",
            value=clean_text(
                editing_record.get("특이사항")
            ),
            placeholder=(
                "예: 제작팀 협의사항, "
                "온열질환 환자 발생 및 조치 내용"
            ),
            height=120,
            key=f"notes_{nonce}",
        )

        save_col, reset_col = st.columns(2)

        with save_col:
            save_clicked = st.button(
                (
                    "수정 내용 저장"
                    if editing_id
                    else "기록 저장"
                ),
                key=f"save_record_{nonce}",
                use_container_width=True,
                type="primary",
                disabled=store_error is not None,
            )

        with reset_col:
            reset_clicked = st.button(
                "입력 초기화",
                key=f"reset_record_{nonce}",
                use_container_width=True,
            )

    if reset_clicked:
        reset_form(go_to_form=True)
        st.rerun()

    if not save_clicked:
        return

    work_start_text = format_time_value(work_start)
    work_end_text = format_time_value(work_end)
    heat_start_text = format_time_value(heat_start)
    heat_end_text = format_time_value(heat_end)
    # 분 단위가 명시된 휴식 조치만 Sheets의 휴게시간 열에 자동 반영합니다.
    rest_minutes = (
        20
        if "2시간 이내 20분 이상 휴식" in measures
        else 10
        if "1시간 이내 10분 이상 휴식" in measures
        else 0
    )
    saved_measures = list(measures)
    if "근무 시간대 조정" in measures and time_adjust_reason.strip():
        saved_measures.append(
            f"{TIME_ADJUST_REASON_PREFIX} {time_adjust_reason.strip()}"
        )
    if other_measure.strip():
        saved_measures.append(other_measure.strip())

    validation_errors: list[str] = []

    if not measures:
        st.warning("예정 또는 시행할 조치를 선택해 주세요")
        return

    if not site.strip():
        validation_errors.append(
            "근무장소를 입력해 주세요."
        )

    if not team:
        validation_errors.append("부서를 선택해 주세요.")

    if worker_count < 1:
        validation_errors.append("근무자 수를 1명 이상 입력해 주세요.")

    if "근무 시간대 조정" in measures and not time_adjust_reason.strip():
        validation_errors.append("근무시간 조정 사유를 입력해 주세요.")

    if not work_start_text or not work_end_text:
        validation_errors.append(
            "근무 시작과 종료 시간을 기록해 주세요."
        )

    if bool(heat_start_text) != bool(heat_end_text):
        validation_errors.append(
            "폭염 노출 시간은 시작과 종료를 "
            "모두 입력하거나 모두 비워 주세요."
        )

    normalized_temperature, temperature_error = (
        normalize_temperature_text(temperature)
    )

    if temperature_error:
        validation_errors.append(temperature_error)

    if validation_errors:
        for message in validation_errors:
            st.error(message)
        return

    now_text = datetime.now(KST).strftime(
        "%Y-%m-%d %H:%M:%S"
    )
    created_at = (
        clean_text(
            editing_record.get("등록시간")
        )
        or now_text
    )
    record_id = editing_id or str(uuid.uuid4())

    record = {
        "id": record_id,
        "작업날짜": work_date.isoformat(),
        "종목": sport,
        "현장명": site.strip(),
        "팀": team,
        "근무자수": str(worker_count),
        "직원": str(employee_count),
        "도급": str(contractor_count),
        "근무시작": work_start_text,
        "근무종료": work_end_text,
        "작성자": author.strip(),
        "폭염시작": heat_start_text,
        "폭염종료": heat_end_text,
        "체감온도": normalized_temperature,
        "휴게시간": str(rest_minutes),
        "공통 조치사항": common_measures,
        "조치사항": " | ".join(saved_measures),
        "특이사항": notes.strip(),
        "등록시간": created_at,
        "수정시간": now_text,
    }

    try:
        if editing_id:
            update_record(editing_id, record)
            st.session_state.flash = (
                "기록을 수정했습니다."
            )
        else:
            append_record(record)
            st.session_state.flash = (
                "사전 예보 기록을 저장했습니다. 작업 후 실제값으로 수정해 주세요."
                if work_date > datetime.now(KST).date()
                else "기록을 저장했습니다."
            )
    except Exception as exc:  # noqa: BLE001
        st.error(f"저장에 실패했습니다: {exc}")
        return

    reset_form(go_to_form=False)
    st.session_state.page = "records"
    st.rerun()


def render_metrics(records: pd.DataFrame) -> None:
    hot_count = 0
    rest_action_count = 0

    if not records.empty:
        hot_count = sum(
            (
                max_temperature(value)
                if max_temperature(value) is not None
                else -999
            )
            >= 33
            for value in records["체감온도"]
        )

        rest_action_count = sum(
            "휴식" in clean_text(value)
            for value in records["조치사항"]
        )

    st.markdown(
        f"""
        <div class="metric-grid">
            <div class="metric-card">
                <span>전체 기록</span>
                <b>{len(records)}</b>
            </div>
            <div class="metric-card">
                <span>체감온도 33℃ 이상</span>
                <b>{hot_count}</b>
            </div>
            <div class="metric-card">
                <span>휴식 조치 기록</span>
                <b>{rest_action_count}</b>
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )


def render_records(
    records: pd.DataFrame,
    store_error: Exception | None,
) -> None:
    st.markdown("## 기록 조회")

    spreadsheet_url = get_secret(
        ("app", "spreadsheet_url"),
        SPREADSHEET_URL_FALLBACK,
    )

    st.link_button(
        "Google Sheets 열기",
        spreadsheet_url,
        use_container_width=True,
    )

    if store_error is not None:
        st.info(
            "Google Sheets 연결이 완료되면 "
            "공동 기록이 여기에 표시됩니다."
        )
        return

    render_metrics(records)

    with st.expander(
        "검색 및 필터",
        expanded=False,
    ):
        search_text = st.text_input(
            "업무내용·근무장소·작성자 검색",
            placeholder="검색어 입력",
            key="record_search",
        )

        team_filter = st.selectbox(
            "부서",
            ["전체"] + TEAM_OPTIONS,
            key="team_filter",
        )

        sport_filter = st.selectbox(
            "업무내용",
            ["전체"] + SPORT_OPTIONS,
            key="sport_filter",
        )

    filtered = records.copy()

    if search_text.strip() and not filtered.empty:
        keyword = search_text.strip().lower()
        mask = (
            filtered["종목"]
            .astype(str)
            .str.lower()
            .str.contains(
                keyword,
                na=False,
                regex=False,
            )
            |
            filtered["현장명"]
            .astype(str)
            .str.lower()
            .str.contains(
                keyword,
                na=False,
                regex=False,
            )
            |
            filtered["작성자"]
            .astype(str)
            .str.lower()
            .str.contains(
                keyword,
                na=False,
                regex=False,
            )
        )
        filtered = filtered[mask]

    if (
        team_filter != "전체"
        and not filtered.empty
    ):
        filtered = filtered[
            filtered["팀"] == team_filter
        ]

    if (
        sport_filter != "전체"
        and not filtered.empty
    ):
        normalized_sports = filtered.apply(
            lambda row: normalize_sport(
                row.get("종목"),
                row.get("현장명"),
            ),
            axis=1,
        )
        filtered = filtered[
            normalized_sports == sport_filter
        ]

    if not filtered.empty:
        filtered = filtered.assign(
            _sort_key=(
                filtered["작업날짜"].astype(str)
                + " "
                + filtered["등록시간"].astype(str)
            )
        ).sort_values(
            "_sort_key",
            ascending=False,
        )

    refresh_col, download_col = st.columns(2)

    with refresh_col:
        if st.button(
            "새로고침",
            use_container_width=True,
            key="refresh_records",
        ):
            load_records.clear()
            st.rerun()

    with download_col:
        st.download_button(
            "엑셀로 추출하기",
            data=make_excel_bytes(records),
            file_name=(
                "현장_폭염_조치_전체기록_"
                f"{datetime.now(KST).strftime('%Y%m%d_%H%M')}"
                ".xlsx"
            ),
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "spreadsheetml.sheet"
            ),
            use_container_width=True,
            disabled=records.empty,
        )

    st.caption(
        "엑셀 파일은 필터와 관계없이 전체 기록을 저장하며, "
        "헤더·테두리·줄바꿈·특이사항 강조가 포함됩니다."
    )

    if filtered.empty:
        st.info("조건에 맞는 기록이 없습니다.")
        return

    st.markdown("### 폭염 보고서 다운로드")
    st.caption(
        "같은 근무일자와 근무장소의 기록을 "
        "Word DOCX 보고서 한 장으로 묶습니다."
    )

    report_source = records.copy()
    if search_text.strip() and not report_source.empty:
        keyword = search_text.strip().lower()
        report_source = report_source[
            report_source["종목"].astype(str).str.lower().str.contains(
                keyword,
                na=False,
                regex=False,
            )
            | report_source["현장명"].astype(str).str.lower().str.contains(
                keyword,
                na=False,
                regex=False,
            )
            | report_source["작성자"].astype(str).str.lower().str.contains(
                keyword,
                na=False,
                regex=False,
            )
        ]
    if sport_filter != "전체" and not report_source.empty:
        report_source = report_source[
            report_source.apply(
                lambda row: normalize_sport(row.get("종목"), row.get("현장명")),
                axis=1,
            )
            == sport_filter
        ]
    report_source = canonicalize_report_rows(report_source)
    report_source["보고서장소"] = report_source["현장명"]
    report_candidates = (
        report_source[["작업날짜", "_보고서장소키", "보고서장소"]]
        .drop_duplicates(subset=["작업날짜", "_보고서장소키"])
        .sort_values(["작업날짜", "보고서장소"], ascending=[False, True])
        .to_dict("records")
    )
    report_labels = [
        f"{clean_text(item['작업날짜'])} · {clean_text(item['보고서장소'])}"
        for item in report_candidates
    ]
    selected_report_label = st.selectbox(
        "보고서 대상",
        report_labels,
        key="report_target",
    )
    selected_report_index = report_labels.index(selected_report_label)
    selected_report = report_candidates[selected_report_index]
    selected_date = clean_text(selected_report["작업날짜"])
    date_report_source = records[records["작업날짜"] == selected_date]
    if sport_filter != "전체" and not date_report_source.empty:
        date_report_source = date_report_source[
            date_report_source.apply(
                lambda row: normalize_sport(row.get("종목"), row.get("현장명")),
                axis=1,
            )
            == sport_filter
        ]
    date_report_source = canonicalize_report_rows(date_report_source)
    report_rows = date_report_source[
        date_report_source["_보고서장소키"]
        == selected_report["_보고서장소키"]
    ]

    try:
        report_filename, report_bytes = report_attachment_for_rows(report_rows)
        st.caption(
            "생성된 Word 보고서는 24시간 캐시에 보관되어 "
            "다운로드와 재전송에 바로 사용됩니다."
        )
        st.download_button(
            "Word 보고서 다운로드",
            data=report_bytes,
            file_name=report_filename,
            mime=(
                "application/vnd.openxmlformats-officedocument."
                "wordprocessingml.document"
            ),
            use_container_width=True,
            key="download_heat_report",
        )
        selected_mail_col, date_mail_col = st.columns(2)
        with selected_mail_col:
            if st.button(
                "선택 보고서 발송",
                use_container_width=True,
                key="send_heat_report_email",
            ):
                try:
                    with st.spinner("보고서를 메일로 발송하고 있습니다..."):
                        recipient_count = send_report_email(
                            clean_text(selected_report["작업날짜"]),
                            clean_text(selected_report["보고서장소"]),
                            report_filename,
                            report_bytes,
                        )
                    st.success(
                        f"보고서를 수신자 {recipient_count}명에게 발송했습니다."
                    )
                except Exception as mail_exc:  # noqa: BLE001
                    st.error(f"메일 발송에 실패했습니다: {mail_exc}")
        with date_mail_col:
            if st.button(
                "해당 날짜 전체 발송",
                use_container_width=True,
                key="send_all_date_reports_email",
            ):
                try:
                    selected_date = clean_text(selected_report["작업날짜"])
                    date_rows = records[records["작업날짜"] == selected_date]
                    date_rows = canonicalize_report_rows(date_rows)
                    attachments: list[tuple[str, bytes]] = []
                    with st.spinner("해당 날짜의 모든 보고서를 만들고 발송하고 있습니다..."):
                        for site_key in unique_texts(
                            date_rows["_보고서장소키"].tolist()
                        ):
                            site_rows = date_rows[
                                date_rows["_보고서장소키"] == site_key
                            ]
                            attachments.append(report_attachment_for_rows(site_rows))
                        recipient_count = send_report_attachments_email(
                            selected_date,
                            "전체 근무장소",
                            attachments,
                        )
                    st.success(
                        f"보고서 {len(attachments)}개를 "
                        f"수신자 {recipient_count}명에게 발송했습니다."
                    )
                except Exception as mail_exc:  # noqa: BLE001
                    st.error(f"날짜 전체 메일 발송에 실패했습니다: {mail_exc}")
    except Exception as exc:  # noqa: BLE001
        st.warning(f"보고서를 만들 수 없습니다: {exc}")

    admin_pin = get_secret(
        ("security", "admin_pin")
    )

    for _, row in filtered.iterrows():
        record = row.to_dict()
        record_id = clean_text(record.get("id"))
        site = (
            markdown_escape(record.get("현장명"))
            or "근무장소 미입력"
        )
        work_date = markdown_escape(
            record.get("작업날짜")
        )
        temperature_text = temperature_display(
            record.get("체감온도")
        )
        temperature_label = (
            f" · {temperature_text}℃"
            if temperature_text
            else ""
        )

        with st.container(border=True):
            status_class = heat_level_class(
                temperature_text
            )
            team_text = (
                clean_text(record.get("팀"))
                or "부서 미입력"
            )
            sport_text = (
                normalize_sport(
                    record.get("종목"),
                    record.get("현장명"),
                )
                or "업무내용 미입력"
            )
            author_text = (
                clean_text(record.get("작성자"))
                or "-"
            )
            worker_text = clean_text(record.get("근무자수")) or "-"
            employee_text = clean_text(record.get("직원")) or "0"
            contractor_text = clean_text(record.get("도급")) or "0"
            work_time_text = (
                f"{clean_text(record.get('근무시작')) or '-'}"
                " ~ "
                f"{clean_text(record.get('근무종료')) or '-'}"
            )
            heat_time_text = (
                f"{clean_text(record.get('폭염시작')) or '-'}"
                " ~ "
                f"{clean_text(record.get('폭염종료')) or '-'}"
            )
            measure_text = clean_text(record.get("조치사항"))
            # 기존 10분 휴식 기록도 조회 화면에서는 계속 표시합니다.
            has_10_minute_rest = (
                "1시간 이내 10분 이상 휴식" in measure_text
            )
            has_20_minute_rest = (
                "2시간 이내 20분 이상 휴식" in measure_text
            )
            if has_10_minute_rest and has_20_minute_rest:
                rest_action_text = "10분·20분 휴식조치"
            elif has_10_minute_rest:
                rest_action_text = "10분 휴식조치"
            elif has_20_minute_rest:
                rest_action_text = "20분 휴식조치"
            else:
                rest_action_text = "-"

            st.markdown(f"### {site}")

            pill_text = (
                f"{heat_level(temperature_text)}"
                f"{temperature_label}"
            )
            st.markdown(
                (
                    f'<span class="status-pill '
                    f'{status_class}">'
                    f'{html.escape(pill_text)}'
                    "</span>"
                ),
                unsafe_allow_html=True,
            )

            st.caption(
                f"{work_date} · {sport_text} · {team_text} · "
                f"근무자 {worker_text}명(직원 {employee_text}·도급 {contractor_text}) · "
                f"작성자 {author_text}"
            )

            st.markdown(
                f"""
                <div class="record-details">
                    <div class="record-detail">
                        <span>근무 시간</span>
                        <b>{html.escape(work_time_text)}</b>
                    </div>
                    <div class="record-detail">
                        <span>폭염 노출</span>
                        <b>{html.escape(heat_time_text)}</b>
                    </div>
                    <div class="record-detail">
                        <span>휴식 조치</span>
                        <b>{html.escape(rest_action_text)}</b>
                    </div>
                </div>
                """,
                unsafe_allow_html=True,
            )

            common_text = clean_text(
                record.get("공통 조치사항")
            )
            if common_text:
                st.markdown("**공통 조치사항**")
                st.markdown(common_text)

            st.write(
                "**조치사항**  "
                f"{clean_text(record.get('조치사항')) or '-'}"
            )

            if clean_text(record.get("특이사항")):
                st.write(
                    "**특이사항**  "
                    f"{clean_text(record.get('특이사항'))}"
                )

            st.caption(
                "등록 "
                f"{clean_text(record.get('등록시간'))}"
                " · 수정 "
                f"{clean_text(record.get('수정시간'))}"
            )

            edit_col, delete_col = st.columns(2)

            with edit_col:
                if st.button(
                    "수정",
                    use_container_width=True,
                    key=f"edit_{record_id}",
                ):
                    st.session_state.editing_id = (
                        record_id
                    )
                    st.session_state.form_nonce += 1
                    st.session_state.page = "form"
                    st.session_state.pending_delete = (
                        None
                    )
                    st.rerun()

            with delete_col:
                if st.button(
                    "삭제",
                    use_container_width=True,
                    key=f"delete_{record_id}",
                ):
                    st.session_state.pending_delete = (
                        record_id
                    )
                    st.rerun()

            if (
                st.session_state.pending_delete
                == record_id
            ):
                st.warning(
                    "이 기록을 삭제할까요? "
                    "삭제 후 복구는 Google Sheets "
                    "변경 기록에서만 가능합니다."
                )

                entered_pin = ""

                if admin_pin:
                    entered_pin = st.text_input(
                        "관리자 삭제 PIN",
                        type="password",
                        key=f"delete_pin_{record_id}",
                    )

                confirm_col, cancel_col = st.columns(2)

                with confirm_col:
                    if st.button(
                        "삭제 확인",
                        type="primary",
                        use_container_width=True,
                        key=f"confirm_{record_id}",
                    ):
                        if (
                            admin_pin
                            and entered_pin != admin_pin
                        ):
                            st.error(
                                "관리자 PIN이 "
                                "올바르지 않습니다."
                            )
                        else:
                            try:
                                delete_record(record_id)
                                st.session_state.pending_delete = (
                                    None
                                )
                                st.session_state.flash = (
                                    "기록을 삭제했습니다."
                                )
                                st.rerun()
                            except Exception as exc:  # noqa: BLE001
                                st.error(
                                    f"삭제에 실패했습니다: {exc}"
                                )

                with cancel_col:
                    if st.button(
                        "취소",
                        use_container_width=True,
                        key=f"cancel_{record_id}",
                    ):
                        st.session_state.pending_delete = (
                            None
                        )
                        st.rerun()


init_state()
render_header()
render_navigation()
show_flash()

records_df = empty_dataframe()
connection_error: Exception | None = None

# 새 기록 첫 화면은 Google Sheets 응답을 기다리지 않고 즉시 표시합니다.
# 기록 목록이 필요하거나 기존 기록을 수정할 때만 캐시된 1회 읽기를 실행합니다.
needs_records = (
    st.session_state.page == "records"
    or bool(clean_text(st.session_state.editing_id))
)

if needs_records:
    try:
        records_df = load_records()
    except Exception as exc:  # noqa: BLE001
        connection_error = exc
        render_setup_error(exc)

if st.session_state.page == "form":
    render_form(
        records_df,
        connection_error,
    )
else:
    render_records(
        records_df,
        connection_error,
    )


